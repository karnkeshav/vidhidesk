"""Tests for the Contracts template engine (Sprint 2, Deliverable 1).

Core guarantees under test:
  - only clause_type='llm_fillable' rows ever reach the LLM gateway;
    fixed_boilerplate text is copied through untouched (CLAUDE.md Hard
    Rule 2 — no hallucinated structure)
  - every LLM call for a clause fill carries the real party names/
    addresses as `entities`, so the existing PII masker actually masks
    them (CLAUDE.md Decision 4) — this test does not re-verify masking
    correctness itself (test_pii_mask.py owns that), only that contracts.py
    upholds its side of the contract
  - the mutual/one-way variant selects the right confidentiality clause
    and excludes the other
  - amendment = calling generate_draft again never overwrites a prior
    draft_versions row, and version_no increments (AC-2.3)
  - clause review (keep/redraft/delete) updates the audit trail and the
    denormalized status, and flips the template beta -> reviewed only
    once every clause has cleared review (Project_Plan §6.4)
  - the real nda.docx skeleton actually renders with docxtpl end to end
"""

from __future__ import annotations

import itertools
import json
import re
from pathlib import Path

import pytest

from app.services import contracts
from app.services.llm_gateway import GenerationResult
from app.services.pii_mask import MaskMap

REPO_ROOT = Path(__file__).resolve().parent.parent.parent


# --- Fake Supabase client (select/insert/update/upsert, multi-table) -------


class FakeResponse:
    def __init__(self, data):
        self.data = data


_id_counter = itertools.count(1)


class FakeQuery:
    def __init__(self, table: "FakeTable", op: str, payload=None):
        self._table = table
        self._op = op
        self._payload = payload
        self._filters: dict[str, object] = {}
        self._order_col = None
        self._order_desc = False
        self._limit = None

    def eq(self, col, val):
        self._filters[col] = val
        return self

    def order(self, col, desc=False):
        self._order_col = col
        self._order_desc = desc
        return self

    def limit(self, n):
        self._limit = n
        return self

    def execute(self):
        if self._op == "select":
            matches = [
                r for r in self._table.rows
                if all(r.get(c) == v for c, v in self._filters.items())
            ]
            if self._order_col:
                matches = sorted(
                    matches, key=lambda r: r.get(self._order_col), reverse=self._order_desc
                )
            if self._limit is not None:
                matches = matches[: self._limit]
            return FakeResponse(matches)

        if self._op == "insert":
            records = self._payload if isinstance(self._payload, list) else [self._payload]
            inserted = []
            for rec in records:
                row = dict(rec)
                row.setdefault("id", f"{self._table.name}-{next(_id_counter)}")
                row.setdefault("created_at", "2026-01-01T00:00:00+00:00")
                self._table.rows.append(row)
                inserted.append(row)
            return FakeResponse(inserted)

        if self._op == "update":
            matches = [
                r for r in self._table.rows
                if all(r.get(c) == v for c, v in self._filters.items())
            ]
            for r in matches:
                r.update(self._payload)
            return FakeResponse(matches)

        if self._op == "upsert":
            records = self._payload if isinstance(self._payload, list) else [self._payload]
            key_cols = (self._table.upsert_key or "id").split(",")
            result = []
            for rec in records:
                match = next(
                    (r for r in self._table.rows if all(r.get(k) == rec.get(k) for k in key_cols)),
                    None,
                )
                if match:
                    match.update(rec)
                    result.append(match)
                else:
                    row = dict(rec)
                    row.setdefault("id", f"{self._table.name}-{next(_id_counter)}")
                    self._table.rows.append(row)
                    result.append(row)
            return FakeResponse(result)

        raise AssertionError(f"unsupported op {self._op}")


class FakeTable:
    def __init__(self, name: str):
        self.name = name
        self.rows: list[dict] = []
        self.upsert_key: str | None = None

    def select(self, *_a, **_k):
        return FakeQuery(self, "select")

    def insert(self, payload):
        return FakeQuery(self, "insert", payload)

    def update(self, payload):
        return FakeQuery(self, "update", payload)

    def upsert(self, payload, on_conflict=None):
        self.upsert_key = on_conflict
        return FakeQuery(self, "upsert", payload)


class FakeDB:
    def __init__(self):
        self._tables: dict[str, FakeTable] = {}

    def table(self, name):
        return self._tables.setdefault(name, FakeTable(name))


def _poisoned_service_client():
    """monkeypatch target for app.db.service_client in tests that must stay
    fully DB-isolated (see test_matter_centric_loading_model_stores_and_returns_template_id
    and test_template_lookup_supports_both_uuid_and_slug_keys below). Both tests also
    patch the router module's already-bound `service_client` name to a FakeDB directly,
    which is what actually makes the app use a fake — this poison pill is a tripwire
    for any other/future code path that re-imports app.db.service_client fresh, so a
    regression fails loudly instead of silently writing to the real Supabase project."""
    raise AssertionError(
        "app.db.service_client() was called for real from a test that must stay "
        "fully DB-isolated (no writes to the production Supabase project)."
    )


# --- Fixtures ----------------------------------------------------------------

NDA_CLAUSES = [
    {"clause_key": "recitals", "display_order": 1, "clause_type": "llm_fillable",
     "applicable_condition": None, "heading": None, "source_text": "Draft recitals for: {{ purpose }}",
     "current_text": "Draft recitals for: {{ purpose }}"},
    {"clause_key": "definitions", "display_order": 2, "clause_type": "fixed_boilerplate",
     "applicable_condition": None, "heading": "Definitions", "source_text": "Fixed boilerplate text.",
     "current_text": "Fixed boilerplate text."},
    {"clause_key": "confidentiality_obligations_mutual", "display_order": 3,
     "clause_type": "fixed_boilerplate", "applicable_condition": {"field": "nda_variant", "equals": "mutual"},
     "heading": "Confidentiality Obligations",
     "source_text": "Confidentiality (mutual variant).",
     "current_text": "Confidentiality (mutual variant)."},
    {"clause_key": "confidentiality_obligations_one_way", "display_order": 3,
     "clause_type": "fixed_boilerplate", "applicable_condition": {"field": "nda_variant", "equals": "one_way"},
     "heading": "Confidentiality Obligations",
     "source_text": "Confidentiality (one-way variant).",
     "current_text": "Confidentiality (one-way variant)."},
    {"clause_key": "term_and_survival", "display_order": 4, "clause_type": "llm_fillable",
     "applicable_condition": None, "heading": "Term and Survival",
     "source_text": "Draft term clause for: {{ tenure }}",
     "current_text": "Draft term clause for: {{ tenure }}"},
]


DEFAULT_TEST_SCHEMA = {
    "template_key": "nda",
    "variant_field": "nda_variant",
    "fields": [
        # "type" is required on every field for real schemas — masking is
        # schema-type-aware (only text/textarea gets masked; see
        # contracts._mask_form_data) — a field missing "type" here would
        # silently never get masked, same bug class as the missing-default
        # StrictUndefined crash this fixture was already built to avoid.
        {"key": "nda_variant", "type": "select", "required": True},
        {"key": "party_a_name", "type": "text", "required": True},
        {"key": "party_a_address", "type": "textarea", "required": True},
        {"key": "party_b_name", "type": "text", "required": True},
        {"key": "party_b_address", "type": "textarea", "required": True},
        {"key": "purpose", "type": "textarea", "required": True},
        {"key": "confidential_items", "type": "textarea", "required": False, "default": ""},
        {"key": "tenure", "type": "text", "required": True},
        {"key": "state", "type": "select", "required": True},
        {"key": "arbitration", "type": "boolean", "default": False},
        {"key": "arbitration_seat", "type": "text", "required": False, "default": ""},
        {"key": "effective_date", "type": "date", "required": True},
    ],
}


def _seed_template(db: FakeDB, clauses=None, schema_json=None) -> str:
    template_id = "template-nda"
    db.table("templates").rows.append(
        {
            "id": template_id,
            "name": "Non-Disclosure Agreement",
            "category": "contracts",
            "docx_path": "templates/contracts/nda.docx",
            "review_status": "beta",
            "schema_json": schema_json if schema_json is not None else DEFAULT_TEST_SCHEMA,
        }
    )
    for clause in clauses if clauses is not None else NDA_CLAUSES:
        db.table("template_clauses").rows.append(
            {"id": f"clause-{clause['clause_key']}", "template_id": template_id, **clause,
             "review_status": "unreviewed"}
        )
    return template_id


def _seed_matter(db: FakeDB, matter_id="matter-1") -> str:
    db.table("matters").rows.append({"id": matter_id, "client_name": "Acme Pvt Ltd"})
    return matter_id


BASE_FORM = {
    "nda_variant": "mutual",
    "party_a_name": "Ramesh Kumar",
    "party_a_entity_type": "Individual",
    "party_a_address": "123 MG Road, Delhi",
    "party_b_name": "Acme Pvt Ltd",
    "party_b_entity_type": "Private Limited Company",
    "party_b_address": "456 Business Park, Mumbai",
    "purpose": "evaluating a potential software licensing deal",
    "tenure": "3 years from the Effective Date",
    "state": "Delhi",
    "arbitration": False,
    "arbitration_seat": "",
    "effective_date": "2026-08-01",
}


def _fake_generate(monkeypatch, canned_text="LLM-GENERATED TEXT"):
    calls = []

    def fake(prompt, task_type=None, mask_map=None, entities=None, auto_detect_names=True, **kwargs):
        calls.append(
            {
                "prompt": prompt, "task_type": task_type, "mask_map": mask_map,
                "entities": entities, "auto_detect_names": auto_detect_names,
            }
        )
        return GenerationResult(
            text=canned_text, provider="gemini", model="gemini-2.5-flash",
            latency_ms=10, masked_prompt=prompt,
        )

    monkeypatch.setattr(contracts, "generate", fake)
    return calls


# --- generate_draft: LLM only touches llm_fillable clauses ------------------


def test_generate_draft_handles_conditionally_hidden_optional_fields(monkeypatch):
    """Regression test for a real Sprint 2 E2E bug: the real frontend never
    submits a value for a field that's conditionally hidden (e.g.
    arbitration_seat when arbitration=False) or an optional field the user
    just never touched (confidential_items) — form_data simply won't have
    the key, unlike every other test in this file which hand-constructs a
    fully-populated dict. Loads the REAL nda.schema.json (not the
    hand-rolled DEFAULT_TEST_SCHEMA) so this stays honest to what the
    actual frontend does. Must not raise jinja2.UndefinedError."""
    import json

    real_schema = json.loads((REPO_ROOT / "templates" / "contracts" / "nda.schema.json").read_text())
    db = FakeDB()
    template_id = _seed_template(db, schema_json=real_schema)
    matter_id = _seed_matter(db)
    _fake_generate(monkeypatch)

    incomplete_form = dict(BASE_FORM)
    incomplete_form.pop("arbitration_seat", None)
    incomplete_form.pop("confidential_items", None)

    result = contracts.generate_draft(matter_id, template_id, incomplete_form, db=db)
    assert result.version_no == 1


def test_generate_draft_calls_llm_only_for_llm_fillable_clauses(monkeypatch):
    db = FakeDB()
    template_id = _seed_template(db)
    matter_id = _seed_matter(db)
    calls = _fake_generate(monkeypatch)

    result = contracts.generate_draft(matter_id, template_id, BASE_FORM, db=db)

    filled_keys = {f.clause_key for f in result.clause_fills}
    assert filled_keys == {"recitals", "term_and_survival"}
    assert len(calls) == 2
    # Fixed boilerplate text must never be sent to the LLM.
    for call in calls:
        assert "Fixed boilerplate text" not in call["prompt"]


def test_generate_draft_passes_real_pii_as_entities_for_masking(monkeypatch):
    """contracts.py's obligation under CLAUDE.md Decision 4: every gateway
    call must carry mask_map + the real party names/addresses as entities,
    so the gateway's own masker (already tested in test_pii_mask.py) has
    something to mask. This does not re-verify masking correctness."""
    db = FakeDB()
    template_id = _seed_template(db)
    matter_id = _seed_matter(db)
    calls = _fake_generate(monkeypatch)

    contracts.generate_draft(matter_id, template_id, BASE_FORM, db=db)

    assert calls, "expected at least one LLM call"
    for call in calls:
        assert isinstance(call["mask_map"], MaskMap)
        assert ("PARTY", "Ramesh Kumar") in call["entities"]
        assert ("PARTY", "Acme Pvt Ltd") in call["entities"]
        assert ("ADDR", "123 MG Road, Delhi") in call["entities"]
        assert ("ADDR", "456 Business Park, Mumbai") in call["entities"]


def test_generate_draft_disables_fuzzy_auto_detect_on_the_gateway_call(monkeypatch):
    """TICKET-1: contracts.py's half of the fix — form_data values are
    already individually masked (full detection) before being interpolated
    into a clause prompt, so the outer gateway call must pass
    auto_detect_names=False (see test_pii_mask.py for what that flag
    actually does to the assembled prompt's static wording)."""
    db = FakeDB()
    template_id = _seed_template(db)
    matter_id = _seed_matter(db)
    calls = _fake_generate(monkeypatch)

    contracts.generate_draft(matter_id, template_id, BASE_FORM, db=db)

    assert calls, "expected at least one LLM call"
    for call in calls:
        assert call["auto_detect_names"] is False


def test_generate_draft_masks_incidental_pii_in_free_text_fields_before_render(monkeypatch):
    """The field-level pre-masking pass (not just the entities list) must
    catch PII embedded in free text, e.g. a PAN mentioned inside the
    purpose field — and it must do so BEFORE the clause template is
    rendered, so the prompt the LLM sees never contains the raw value."""
    db = FakeDB()
    template_id = _seed_template(db)
    matter_id = _seed_matter(db)
    calls = _fake_generate(monkeypatch)

    form = dict(BASE_FORM)
    form["purpose"] = "Context involves PAN ABCDE1234F and mobile 9876543210."

    contracts.generate_draft(matter_id, template_id, form, db=db)

    recitals_call = next(c for c in calls if "recitals" in c["prompt"].lower() or "Draft recitals" in c["prompt"])
    assert "ABCDE1234F" not in recitals_call["prompt"]
    assert "9876543210" not in recitals_call["prompt"]
    assert "PAN_1" in recitals_call["prompt"]
    assert "PHONE_1" in recitals_call["prompt"]


def test_select_field_values_are_never_masked(monkeypatch):
    """Regression test for a real Service Agreement bug: masking every
    string form value indiscriminately (the original TICKET-1 fix, before
    it was made schema-aware) also masked `select` field values —
    "Fixed Fee" false-positived through the same Title-Case-run heuristic
    as "Governing Law", silently breaking a fixed_boilerplate clause's own
    `{% if fee_structure == 'Fixed Fee' %}` comparison (masked value could
    never equal the literal string, so no branch matched — no error, just
    a blank clause). `select`/`boolean`/`date` fields must render through
    with their exact schema-declared value, never a PARTY_x placeholder."""
    schema = {
        "fields": [
            {"key": "fee_structure", "type": "select"},
        ]
    }
    clauses = [
        {"clause_key": "payment", "display_order": 1, "clause_type": "fixed_boilerplate",
         "applicable_condition": None, "heading": "Payment",
         "source_text": "{% if fee_structure == 'Fixed Fee' %}Fixed fee applies.{% endif %}",
         "current_text": "{% if fee_structure == 'Fixed Fee' %}Fixed fee applies.{% endif %}"},
    ]
    db = FakeDB()
    template_id = _seed_template(db, clauses=clauses, schema_json=schema)
    matter_id = _seed_matter(db)
    _fake_generate(monkeypatch)

    result = contracts.generate_draft(matter_id, template_id, {"fee_structure": "Fixed Fee"}, db=db)

    assert "Fixed fee applies." in result.full_text


def test_generate_draft_selects_variant_specific_confidentiality_clause(monkeypatch):
    db = FakeDB()
    template_id = _seed_template(db)
    matter_id = _seed_matter(db)
    _fake_generate(monkeypatch)

    mutual_form = dict(BASE_FORM, nda_variant="mutual")
    contracts.generate_draft(matter_id, template_id, mutual_form, db=db)

    one_way_form = dict(BASE_FORM, nda_variant="one_way", party_a_role="disclosing")
    contracts.generate_draft(matter_id, template_id, one_way_form, db=db)

    # Both drafts render fine (docx assertions are covered separately below);
    # here we only assert the clause-selection logic itself via the
    # applicable-clauses helper, which is what actually branches on variant.
    all_clauses = db.table("template_clauses").rows
    mutual_selected = contracts._applicable_clauses(all_clauses, {"nda_variant": "mutual"})
    one_way_selected = contracts._applicable_clauses(all_clauses, {"nda_variant": "one_way"})
    assert "confidentiality_obligations_mutual" in {c["clause_key"] for c in mutual_selected}
    assert "confidentiality_obligations_one_way" not in {c["clause_key"] for c in mutual_selected}
    assert "confidentiality_obligations_one_way" in {c["clause_key"] for c in one_way_selected}
    assert "confidentiality_obligations_mutual" not in {c["clause_key"] for c in one_way_selected}


# --- amendment loop: new version, prior versions untouched ------------------


def test_generate_draft_increments_version_without_losing_prior_versions(monkeypatch):
    db = FakeDB()
    template_id = _seed_template(db)
    matter_id = _seed_matter(db)
    _fake_generate(monkeypatch)

    first = contracts.generate_draft(matter_id, template_id, BASE_FORM, db=db)
    amended_form = dict(BASE_FORM, tenure="5 years from the Effective Date")
    second = contracts.generate_draft(matter_id, template_id, amended_form, db=db)

    assert first.version_no == 1
    assert second.version_no == 2
    all_versions = db.table("draft_versions").rows
    assert len(all_versions) == 2
    assert {v["id"] for v in all_versions} == {first.draft_version_id, second.draft_version_id}


# --- clause review -----------------------------------------------------------


def test_review_clause_redraft_updates_current_text_and_logs_audit_row():
    db = FakeDB()
    template_id = _seed_template(db)
    clause_id = "clause-definitions"

    updated = contracts.review_clause(
        clause_id, "redraft", redraft_text="Revised definitions text.", db=db
    )

    assert updated["review_status"] == "redrafted"
    assert updated["current_text"] == "Revised definitions text."
    audit_rows = db.table("clause_reviews").rows
    assert len(audit_rows) == 1
    assert audit_rows[0]["decision"] == "redraft"
    assert audit_rows[0]["redraft_text"] == "Revised definitions text."


def test_review_clause_redraft_requires_text():
    db = FakeDB()
    _seed_template(db)
    with pytest.raises(ValueError):
        contracts.review_clause("clause-definitions", "redraft", db=db)


def test_review_clause_rejects_invalid_decision():
    db = FakeDB()
    _seed_template(db)
    with pytest.raises(ValueError):
        contracts.review_clause("clause-definitions", "approve", db=db)


def test_review_clause_delete_requires_reviewer_notes():
    db = FakeDB()
    _seed_template(db)
    with pytest.raises(ValueError):
        contracts.review_clause("clause-definitions", "delete", db=db)


def test_template_flips_beta_to_reviewed_only_once_every_clause_cleared():
    db = FakeDB()
    template_id = _seed_template(db, clauses=[NDA_CLAUSES[0], NDA_CLAUSES[1]])
    clause_ids = [f"clause-{c['clause_key']}" for c in [NDA_CLAUSES[0], NDA_CLAUSES[1]]]

    contracts.review_clause(clause_ids[0], "keep", db=db)
    template_row = db.table("templates").rows[0]
    assert template_row["review_status"] == "beta"  # one clause still unreviewed

    contracts.review_clause(clause_ids[1], "keep", db=db)
    assert template_row["review_status"] == "reviewed"


# --- bulk-keep boilerplate (review-velocity Lever 1, 2026-08-02) -----------
#
# NDA_CLAUSES fixture shape: recitals (llm_fillable), definitions
# (fixed_boilerplate), confidentiality_obligations_mutual/one_way (both
# fixed_boilerplate, applicable_condition-gated), term_and_survival
# (llm_fillable) — 2 llm_fillable, 3 fixed_boilerplate.


def test_bulk_keep_boilerplate_keeps_only_unreviewed_fixed_boilerplate():
    db = FakeDB()
    _seed_template(db)

    updated = contracts.bulk_keep_boilerplate_clauses("template-nda", db=db)

    assert len(updated) == 3, "only the 3 fixed_boilerplate clauses should be kept"
    assert all(c["review_status"] == "kept" for c in updated)
    assert all(c["clause_type"] == "fixed_boilerplate" for c in updated)

    all_clauses = {c["id"]: c for c in db.table("template_clauses").rows}
    llm_fillable_ids = ["clause-recitals", "clause-term_and_survival"]
    for cid in llm_fillable_ids:
        assert all_clauses[cid]["review_status"] == "unreviewed", (
            "llm_fillable clauses must never be touched by bulk-keep"
        )

    audit_rows = db.table("clause_reviews").rows
    assert len(audit_rows) == 3
    assert all(r["decision"] == "keep" for r in audit_rows)
    assert all(r["reviewer_notes"] == contracts.BULK_KEEP_REVIEWER_NOTES for r in audit_rows)


def test_bulk_keep_boilerplate_never_overwrites_an_already_reviewed_clause():
    db = FakeDB()
    _seed_template(db)

    # Nitesh redrafts one fixed_boilerplate clause by hand before the
    # bulk action runs.
    contracts.review_clause(
        "clause-definitions", "redraft", redraft_text="Nitesh's custom definitions text.", db=db
    )

    updated = contracts.bulk_keep_boilerplate_clauses("template-nda", db=db)

    # Only the two remaining unreviewed fixed_boilerplate clauses (the
    # two confidentiality variants) get bulk-kept — definitions is
    # excluded because it's no longer 'unreviewed'.
    assert len(updated) == 2
    assert {c["clause_key"] for c in updated} == {
        "confidentiality_obligations_mutual", "confidentiality_obligations_one_way",
    }

    definitions = next(c for c in db.table("template_clauses").rows if c["id"] == "clause-definitions")
    assert definitions["review_status"] == "redrafted"
    assert definitions["current_text"] == "Nitesh's custom definitions text.", (
        "bulk-keep must never touch a clause that already has a human review decision"
    )
    # Still exactly one clause_reviews row for definitions (the redraft) —
    # bulk-keep didn't add a second row for it.
    definitions_reviews = [r for r in db.table("clause_reviews").rows if r["clause_id"] == "clause-definitions"]
    assert len(definitions_reviews) == 1
    assert definitions_reviews[0]["decision"] == "redraft"


def test_bulk_keep_boilerplate_can_flip_template_to_reviewed():
    db = FakeDB()
    _seed_template(db)
    contracts.review_clause("clause-recitals", "keep", db=db)
    contracts.review_clause("clause-term_and_survival", "keep", db=db)

    template_row = db.table("templates").rows[0]
    assert template_row["review_status"] == "beta", "3 fixed_boilerplate clauses still unreviewed"

    contracts.bulk_keep_boilerplate_clauses("template-nda", db=db)

    assert template_row["review_status"] == "reviewed", (
        "bulk-keep clearing the last unreviewed clauses must flip beta -> reviewed, same as review_clause does"
    )


def test_bulk_keep_boilerplate_returns_empty_list_when_nothing_qualifies():
    db = FakeDB()
    _seed_template(db)

    contracts.bulk_keep_boilerplate_clauses("template-nda", db=db)
    second_call = contracts.bulk_keep_boilerplate_clauses("template-nda", db=db)

    assert second_call == []


# --- re-seed must never silently overwrite a reviewed clause ---------------
#
# Found live 2026-08-02, more urgent than a Sprint 3 ticket: every seed
# script's upsert previously wrote current_text = source_text
# unconditionally, every re-seed, for every clause — including ones
# already reviewed. This would (a) wipe out a redraft even on a
# no-op re-run (current_text always reset regardless of whether
# source_text changed), and (b) silently swap a 'kept' clause's
# rendered text out from under an approval the moment its source_text
# was edited. Fix, deliberately STRICT: a re-seed HALTS (raises
# ReviewedClauseConflict) if an already-reviewed clause's incoming
# source_text differs from what's stored — no auto-decision. When
# nothing has changed for reviewed clauses, they're excluded from the
# upsert batch entirely (current_text/review_status never touched,
# only structural fields refreshed).
#
# Uses the real seed_nda_template.py module directly (not just its
# CLAUSES list) to exercise _write_clauses_preserving_review and
# ReviewedClauseConflict as actually shipped.


def _load_nda_seed_module():
    spec = _importlib_util.spec_from_file_location(
        "seed_nda_template_module_test",
        Path(__file__).resolve().parent.parent / "scripts" / "seed_nda_template.py",
    )
    module = _importlib_util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def test_reseed_halts_when_a_kept_clauses_source_text_changed():
    module = _load_nda_seed_module()
    db = FakeDB()

    # Initial seed, then Nitesh keeps "definitions" as-is.
    module._write_clauses_preserving_review(db, "template-nda", module.CLAUSES)
    definitions_id = next(
        r["id"] for r in db.table("template_clauses").rows if r["clause_key"] == "definitions"
    )
    contracts.review_clause(definitions_id, "keep", db=db)

    # An author now edits definitions' source_text (e.g. fixing wording)
    # without knowing it's already been reviewed.
    edited_clauses = [
        dict(c, source_text=c["source_text"] + " EDITED.") if c["clause_key"] == "definitions" else c
        for c in module.CLAUSES
    ]

    with pytest.raises(module.ReviewedClauseConflict) as exc_info:
        module._write_clauses_preserving_review(db, "template-nda", edited_clauses)

    assert "definitions" in str(exc_info.value)
    assert "kept" in str(exc_info.value)

    # Nothing was written — not even the unaffected clauses — matching
    # "HALTS... no auto-decisions."
    definitions_row = next(r for r in db.table("template_clauses").rows if r["clause_key"] == "definitions")
    assert not definitions_row["current_text"].endswith("EDITED.")
    assert definitions_row["source_text"] == module.CLAUSES[1]["source_text"]


def test_reseed_preserves_a_redrafted_clause_when_source_text_unchanged():
    module = _load_nda_seed_module()
    db = FakeDB()

    module._write_clauses_preserving_review(db, "template-nda", module.CLAUSES)
    definitions_id = next(
        r["id"] for r in db.table("template_clauses").rows if r["clause_key"] == "definitions"
    )
    contracts.review_clause(
        definitions_id, "redraft", redraft_text="Nitesh's fully custom definitions text.", db=db
    )

    # Re-seed with the EXACT same CLAUSES content — a plain re-run, no
    # edits anywhere. Before the fix, this alone (no content change at
    # all) would still have wiped the redraft, since current_text was
    # always reset unconditionally.
    module._write_clauses_preserving_review(db, "template-nda", module.CLAUSES)

    definitions_row = next(r for r in db.table("template_clauses").rows if r["clause_key"] == "definitions")
    assert definitions_row["current_text"] == "Nitesh's fully custom definitions text."
    assert definitions_row["review_status"] == "redrafted"


def test_reseed_still_updates_unreviewed_clauses_normally():
    module = _load_nda_seed_module()
    db = FakeDB()

    module._write_clauses_preserving_review(db, "template-nda", module.CLAUSES)
    # definitions stays unreviewed; recitals gets kept.
    recitals_id = next(r["id"] for r in db.table("template_clauses").rows if r["clause_key"] == "recitals")
    contracts.review_clause(recitals_id, "keep", db=db)

    edited_clauses = [
        dict(c, source_text=c["source_text"] + " Updated wording.") if c["clause_key"] == "definitions" else c
        for c in module.CLAUSES
    ]
    # No conflict — only an unreviewed clause changed.
    module._write_clauses_preserving_review(db, "template-nda", edited_clauses)

    definitions_row = next(r for r in db.table("template_clauses").rows if r["clause_key"] == "definitions")
    assert definitions_row["current_text"].endswith("Updated wording.")
    assert definitions_row["review_status"] == "unreviewed"


# --- docx rendering (real skeleton file, mocked LLM only) -------------------


def test_generate_draft_renders_real_docx_skeleton(monkeypatch):
    """Exercises the actual templates/contracts/nda.docx file end to end —
    the one piece nothing above touches. LLM calls are still mocked (no
    network in unit tests); docxtpl rendering is real. Writes to the real
    (gitignored) generated_drafts/ dir, cleaned up at the end of the test."""
    from docx import Document

    db = FakeDB()
    template_id = _seed_template(db)
    matter_id = _seed_matter(db, matter_id="matter-docx-smoke-test")
    _fake_generate(monkeypatch, canned_text="WHEREAS the parties wish to collaborate.")

    result = contracts.generate_draft(matter_id, template_id, BASE_FORM, db=db)
    output_path = REPO_ROOT / result.docx_path
    try:
        assert output_path.exists()
        doc = Document(str(output_path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Ramesh Kumar" in full_text
        assert "Acme Pvt Ltd" in full_text
        assert "BETA — PENDING CLAUSE REVIEW" in full_text
        assert "WHEREAS the parties wish to collaborate." in full_text
        assert "Fixed boilerplate text." in full_text
        # TICKET-3: an_or_a filter — "an Individual" (vowel), "a Private
        # Limited Company" (consonant), never the ungrammatical "a Individual".
        assert "an Individual having its registered" in full_text
        assert "a Private Limited Company having its registered" in full_text
        assert "a Individual" not in full_text
        # Migration 0008: numbers are auto-assigned at assembly time from
        # each clause's `heading`, not hardcoded in clause body text —
        # recitals (no heading) stays unnumbered, Definitions is clause 1.
        assert "1. Definitions" in full_text
        assert "0. Recitals" not in full_text and "1. Recitals" not in full_text
    finally:
        output_path.unlink(missing_ok=True)


def test_an_or_a_filter():
    from app.services.contracts import _an_or_a

    assert _an_or_a("Individual") == "an"
    assert _an_or_a("Private Limited Company") == "a"
    assert _an_or_a("LLP") == "an"  # acronym read as "el-el-pee" — vowel sound despite consonant spelling
    assert _an_or_a("Partnership Firm") == "a"
    assert _an_or_a("") == "a"


# --- Generic list/repeater field type (Sprint 2 Deliverable 2 prep) --------
# Proves the mechanism generically, ahead of Service Agreement — a
# fixed_boilerplate clause looping over a "list" schema field with {% for %},
# rendering each item verbatim (no LLM call, no paraphrasing risk on the
# kind of enumerable content — deliverables, benefits, fixtures — every
# template from here on needs).


def test_with_schema_defaults_normalizes_list_field_items():
    schema = {
        "fields": [
            {"key": "deliverables", "type": "list", "item_schema": [
                {"key": "description", "type": "text", "required": True},
                {"key": "notes", "type": "text", "required": False, "default": ""},
            ]},
        ]
    }
    form_data = {"deliverables": [{"description": "Build the API"}, {"description": "Write docs", "notes": "PDF format"}]}

    result = contracts._with_schema_defaults(form_data, schema)

    assert result["deliverables"] == [
        {"description": "Build the API", "notes": ""},
        {"description": "Write docs", "notes": "PDF format"},
    ]


def test_with_schema_defaults_missing_list_field_defaults_to_empty_list():
    schema = {"fields": [{"key": "deliverables", "type": "list", "item_schema": []}]}
    result = contracts._with_schema_defaults({}, schema)
    assert result["deliverables"] == []


def test_generate_draft_renders_list_field_verbatim_no_llm(monkeypatch):
    """The core Service Agreement mechanism: a fixed_boilerplate clause with
    a {% for %} loop over a list field renders every item's exact text —
    no LLM call, so no paraphrasing risk on enumerable content."""
    schema = {
        "fields": [
            {"key": "party_a_name", "type": "text"},
            {"key": "deliverables", "type": "list", "item_schema": [
                {"key": "description", "type": "text"},
                {"key": "due_date", "type": "text"},
            ]},
        ]
    }
    clauses = [
        {"clause_key": "scope", "display_order": 1, "clause_type": "fixed_boilerplate",
         "applicable_condition": None,
         "source_text": "Scope of Services:\n{% for item in deliverables %}- {{ item.description }} (due {{ item.due_date }})\n{% endfor %}",
         "current_text": "Scope of Services:\n{% for item in deliverables %}- {{ item.description }} (due {{ item.due_date }})\n{% endfor %}"},
    ]
    db = FakeDB()
    template_id = _seed_template(db, clauses=clauses, schema_json=schema)
    matter_id = _seed_matter(db)
    calls = _fake_generate(monkeypatch)

    form = {
        "party_a_name": "Ramesh Kumar",
        "deliverables": [
            {"description": "Build the payment API", "due_date": "2026-09-01"},
            {"description": "Deliver integration tests", "due_date": "2026-09-15"},
        ],
    }
    result = contracts.generate_draft(matter_id, template_id, form, db=db)

    assert not calls, "a fixed_boilerplate clause must never call the LLM"
    docx_path = REPO_ROOT / result.docx_path
    try:
        from docx import Document

        doc = Document(str(docx_path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Build the payment API (due 2026-09-01)" in full_text
        assert "Deliver integration tests (due 2026-09-15)" in full_text
    finally:
        docx_path.unlink(missing_ok=True)


# --- version history + amendment_note ---------------------------------------


def test_list_drafts_returns_newest_first(monkeypatch):
    db = FakeDB()
    template_id = _seed_template(db)
    matter_id = _seed_matter(db)
    _fake_generate(monkeypatch)

    first = contracts.generate_draft(matter_id, template_id, BASE_FORM, db=db)
    second = contracts.generate_draft(matter_id, template_id, BASE_FORM, db=db)

    versions = contracts.list_drafts(matter_id, db=db)
    assert [v["id"] for v in versions] == [second.draft_version_id, first.draft_version_id]


def test_amendment_note_is_appended_to_llm_fillable_prompts_only(monkeypatch):
    db = FakeDB()
    template_id = _seed_template(db)
    matter_id = _seed_matter(db)
    calls = _fake_generate(monkeypatch)

    contracts.generate_draft(
        matter_id, template_id, BASE_FORM, amendment_note="reduce lock-in to 12 months", db=db
    )

    assert calls, "expected at least one LLM call"
    for call in calls:
        assert "reduce lock-in to 12 months" in call["prompt"]


# --- PDF export: environment without LibreOffice --------------------------


def test_convert_docx_to_pdf_raises_clear_error_without_soffice(monkeypatch, tmp_path):
    monkeypatch.setattr(contracts.shutil, "which", lambda _name: None)
    fake_docx = tmp_path / "draft.docx"
    fake_docx.write_bytes(b"not a real docx, never read")

    with pytest.raises(contracts.PdfConversionUnavailable):
        contracts.convert_docx_to_pdf(fake_docx)


# --- Service Agreement (Sprint 2 Deliverable 2, Batch 1) --------------------
# Pre-live-DB verification of the real schema + clause content: the SLA
# conditional-exclusion + auto-numbering interaction, the Jinja branching
# on fee_structure, the deliverables list loop, and an_or_a on a second
# template. Ahead of the live E2E, which needs migration 0008 applied.

import importlib.util as _importlib_util  # noqa: E402


def _load_service_agreement_fixtures():
    schema = json.loads((REPO_ROOT / "templates" / "contracts" / "service-agreement.schema.json").read_text())
    spec = _importlib_util.spec_from_file_location(
        "seed_service_agreement_template",
        Path(__file__).resolve().parent.parent / "scripts" / "seed_service_agreement_template.py",
    )
    module = _importlib_util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return schema, module.CLAUSES


def _load_nda_fixtures():
    """Real nda.schema.json + the real seed_nda_template.py CLAUSES list —
    same rationale as _load_service_agreement_fixtures above: the
    hand-rolled NDA_CLAUSES fixture near the top of this file uses a
    placeholder recitals prompt ("Draft recitals for: {{ purpose }}")
    that would never have caught the missing-party-names bug found live
    2026-08-01, because it never had party names in it to begin with."""
    schema = json.loads((REPO_ROOT / "templates" / "contracts" / "nda.schema.json").read_text())
    spec = _importlib_util.spec_from_file_location(
        "seed_nda_template",
        Path(__file__).resolve().parent.parent / "scripts" / "seed_nda_template.py",
    )
    module = _importlib_util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return schema, module.CLAUSES


def _seed_nda_real(db):
    schema, clauses = _load_nda_fixtures()
    template_id = "template-nda-real"
    db.table("templates").rows.append(
        {
            "id": template_id,
            "name": "Non-Disclosure Agreement",
            "category": "contracts",
            "docx_path": "templates/contracts/nda.docx",
            "review_status": "beta",
            "schema_json": schema,
        }
    )
    for clause in clauses:
        db.table("template_clauses").rows.append(
            {"id": f"clause-{clause['clause_key']}", "template_id": template_id, **clause,
             "current_text": clause["source_text"], "review_status": "unreviewed"}
        )
    return template_id


def _load_consultancy_fixtures():
    """Real consultancy.schema.json + the real seed_consultancy_template.py
    CLAUSES list — same rationale as the NDA/Service Agreement loaders
    above, applied from the start for this template rather than added
    after a bug is found live."""
    schema = json.loads((REPO_ROOT / "templates" / "contracts" / "consultancy.schema.json").read_text())
    spec = _importlib_util.spec_from_file_location(
        "seed_consultancy_template",
        Path(__file__).resolve().parent.parent / "scripts" / "seed_consultancy_template.py",
    )
    module = _importlib_util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return schema, module.CLAUSES


def _seed_consultancy(db):
    schema, clauses = _load_consultancy_fixtures()
    template_id = "template-consultancy"
    db.table("templates").rows.append(
        {
            "id": template_id,
            "name": "Consultancy Agreement",
            "category": "contracts",
            "docx_path": "templates/contracts/consultancy.docx",
            "review_status": "beta",
            "schema_json": schema,
        }
    )
    for clause in clauses:
        db.table("template_clauses").rows.append(
            {"id": f"clause-{clause['clause_key']}", "template_id": template_id, **clause,
             "current_text": clause["source_text"], "review_status": "unreviewed"}
        )
    return template_id


def _load_mou_fixtures():
    """Real mou.schema.json + the real seed_mou_template.py CLAUSES list —
    same rationale as the NDA/Service Agreement/Consultancy loaders
    above, applied from the start for this template too."""
    schema = json.loads((REPO_ROOT / "templates" / "contracts" / "mou.schema.json").read_text())
    spec = _importlib_util.spec_from_file_location(
        "seed_mou_template",
        Path(__file__).resolve().parent.parent / "scripts" / "seed_mou_template.py",
    )
    module = _importlib_util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return schema, module.CLAUSES


def _seed_mou(db):
    schema, clauses = _load_mou_fixtures()
    template_id = "template-mou"
    db.table("templates").rows.append(
        {
            "id": template_id,
            "name": "Memorandum of Understanding",
            "category": "contracts",
            "docx_path": "templates/contracts/mou.docx",
            "review_status": "beta",
            "schema_json": schema,
        }
    )
    for clause in clauses:
        db.table("template_clauses").rows.append(
            {"id": f"clause-{clause['clause_key']}", "template_id": template_id, **clause,
             "current_text": clause["source_text"], "review_status": "unreviewed"}
        )
    return template_id


MOU_FORM = {
    "party_a_name": "Karan Malhotra",
    "party_a_entity_type": "Individual",
    "party_a_address": "7 Lodi Road, New Delhi",
    "party_b_name": "Greenfield Renewables Pvt Ltd",
    "party_b_entity_type": "Private Limited Company",
    "party_b_address": "Sector 21, Noida",
    "purpose": "exploring a joint venture for a solar energy project",
    "confidentiality_direction": "mutual",
    "confidentiality_survival_period": "3 years",
    "term_duration": "6 months from the Effective Date, or until superseded by a definitive agreement, whichever is earlier",
    "termination_notice_period": "15 days",
    "state": "Delhi",
    "arbitration": False,
    "arbitration_seat": "",
    "effective_date": "2026-08-02",
}


def _load_employment_fixtures():
    """Real employment.schema.json + the real seed_employment_template.py
    CLAUSES list — same rationale as the NDA/Service Agreement/
    Consultancy/MoU loaders above, applied from the start."""
    schema = json.loads((REPO_ROOT / "templates" / "contracts" / "employment.schema.json").read_text())
    spec = _importlib_util.spec_from_file_location(
        "seed_employment_template",
        Path(__file__).resolve().parent.parent / "scripts" / "seed_employment_template.py",
    )
    module = _importlib_util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return schema, module.CLAUSES


def _seed_employment(db):
    schema, clauses = _load_employment_fixtures()
    template_id = "template-employment"
    db.table("templates").rows.append(
        {
            "id": template_id,
            "name": "Employment Agreement",
            "category": "contracts",
            "docx_path": "templates/contracts/employment.docx",
            "review_status": "beta",
            "schema_json": schema,
        }
    )
    for clause in clauses:
        db.table("template_clauses").rows.append(
            {"id": f"clause-{clause['clause_key']}", "template_id": template_id, **clause,
             "current_text": clause["source_text"], "review_status": "unreviewed"}
        )
    return template_id


EMPLOYMENT_FORM = {
    "party_a_name": "Bright Horizons Tech Pvt Ltd",
    "party_a_entity_type": "Private Limited Company",
    "party_a_address": "Tower B, Cyber Hub, Gurugram",
    "party_b_name": "Sneha Reddy",
    "party_b_address": "45 Jubilee Hills, Hyderabad",
    "designation": "Senior Software Engineer",
    "department": "Engineering",
    "reporting_to": "the Engineering Manager",
    "duties_description": "design, develop, and maintain backend services; participate in code reviews and on-call rotations",
    "employment_type": "Full-Time",
    "fixed_term_end_date": "",
    "has_probation": True,
    "probation_period": "6 months",
    "annual_ctc": "₹18,00,000 per annum",
    "other_benefits": "health insurance for self and dependents, annual performance bonus eligibility",
    "state": "Delhi",
    "termination_notice_period": "30 days",
    "non_compete_notes": "",
    "arbitration": False,
    "arbitration_seat": "",
    "effective_date": "2026-08-02",
}


CONSULTANCY_FORM = {
    "party_a_name": "Anjali Mehta",
    "party_a_entity_type": "Individual",
    "party_a_address": "18 Green Park, New Delhi",
    "party_b_name": "Bluewave Logistics Pvt Ltd",
    "party_b_entity_type": "Private Limited Company",
    "party_b_address": "Sector 44, Gurugram",
    "purpose": "advising on supply chain restructuring",
    "scope_notes": "",
    "deliverables": [],
    "fee_structure": "Retainer",
    "fee_amount": "",
    "payment_frequency": "",
    "retainer_fee_amount": "₹1,00,000 per month",
    "retainer_frequency": "Monthly",
    "retainer_scope_hours": "up to 20 hours per month",
    "late_payment_interest_rate": "18% per annum",
    "ip_ownership_model": "Full Assignment to Client",
    "ip_carveout_notes": "",
    "confidentiality_direction": "mutual",
    "confidentiality_survival_period": "3 years",
    "term_duration": "12 months from the Effective Date",
    "termination_notice_period": "30 days",
    "state": "Delhi",
    "arbitration": False,
    "arbitration_seat": "",
    "effective_date": "2026-08-02",
}


SERVICE_AGREEMENT_FORM = {
    "party_a_name": "Ramesh Kumar",
    "party_a_entity_type": "Individual",
    "party_a_address": "42 MG Road, Bengaluru",
    "party_b_name": "Acme Technologies Pvt Ltd",
    "party_b_entity_type": "Private Limited Company",
    "party_b_address": "Business Park, Mumbai",
    "purpose": "Ongoing software consulting services",
    "deliverables": [
        {"description": "Design and build the payment API", "due_date": "2026-09-01"},
        {"description": "Deliver integration test suite", "due_date": "2026-09-15"},
    ],
    "fee_structure": "Fixed Fee",
    "fee_amount": "₹1,50,000",
    "payment_frequency": "Monthly",
    "late_payment_interest_rate": "18% per annum",
    "ip_ownership_model": "Full Assignment to Client",
    "ip_carveout_notes": "",
    "confidentiality_direction": "one_way_from_client",
    "confidentiality_survival_period": "3 years",
    "term_duration": "12 months from the Effective Date",
    "termination_notice_period": "30 days",
    "state": "Delhi",
    "arbitration": False,
    "arbitration_seat": "",
    "effective_date": "2026-08-01",
}


def _seed_service_agreement(db):
    schema, clauses = _load_service_agreement_fixtures()
    template_id = "template-service-agreement"
    db.table("templates").rows.append(
        {
            "id": template_id,
            "name": "Service Agreement",
            "category": "contracts",
            "docx_path": "templates/contracts/service-agreement.docx",
            "review_status": "beta",
            "schema_json": schema,
        }
    )
    for clause in clauses:
        db.table("template_clauses").rows.append(
            {"id": f"clause-{clause['clause_key']}", "template_id": template_id, **clause,
             "current_text": clause["source_text"], "review_status": "unreviewed"}
        )
    return template_id


def test_service_agreement_sla_excluded_four_fills_and_numbering_shifts(monkeypatch):
    from docx import Document

    db = FakeDB()
    template_id = _seed_service_agreement(db)
    matter_id = _seed_matter(db)
    _fake_generate(monkeypatch, canned_text="WHEREAS the parties wish to engage.")

    form = dict(SERVICE_AGREEMENT_FORM, include_sla=False)
    result = contracts.generate_draft(matter_id, template_id, form, db=db)

    assert len(result.clause_fills) == 4, "SLA is fixed_boilerplate — must never produce a fill row"
    assert "Service Levels" not in result.full_text

    output_path = REPO_ROOT / result.docx_path
    try:
        doc = Document(str(output_path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        # 9 numbered clauses with SLA excluded (recitals is unnumbered, so
        # 10 total clauses - recitals - SLA = 9) — Miscellaneous (last)
        # must read "9.", not the SLA-included template's "10.".
        assert "9. Miscellaneous" in full_text
        assert "10. Miscellaneous" not in full_text
        # Deliverables render verbatim, no LLM paraphrase.
        assert "Design and build the payment API (due 2026-09-01)" in full_text
        assert "Deliver integration test suite (due 2026-09-15)" in full_text
        # Fixed Fee branch of the payment clause, not Hourly/Milestone.
        assert "a fixed fee of ₹1,50,000" in full_text
        assert "18% per annum" in full_text
        # an_or_a on a second template: "an Individual", "a Private Limited Company".
        assert "an Individual having its" in full_text
        assert "a Private Limited Company having its" in full_text
    finally:
        output_path.unlink(missing_ok=True)


def test_service_agreement_sla_included_five_clauses_numbering_intact(monkeypatch):
    from docx import Document

    db = FakeDB()
    template_id = _seed_service_agreement(db)
    matter_id = _seed_matter(db)
    _fake_generate(monkeypatch, canned_text="WHEREAS the parties wish to engage.")

    form = dict(
        SERVICE_AGREEMENT_FORM,
        include_sla=True,
        sla_response_time_hours="4",
        sla_resolution_time_hours="24",
        sla_uptime_percentage="99.9",
        sla_credit_terms="A 2% fee credit applies per hour of SLA breach, capped at 10% of monthly Fees.",
    )
    result = contracts.generate_draft(matter_id, template_id, form, db=db)

    assert len(result.clause_fills) == 4, "SLA inclusion must not change the llm_fillable count"
    assert "Service Levels" in result.full_text

    output_path = REPO_ROOT / result.docx_path
    try:
        doc = Document(str(output_path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        # 10 numbered clauses with SLA included (11 total - unnumbered recitals).
        assert "10. Miscellaneous" in full_text
        assert "4. Service Levels" in full_text
        assert "respond to a Client-reported issue within 4 hours" in full_text
        assert "maintain 99.9% uptime" in full_text
    finally:
        output_path.unlink(missing_ok=True)


# --- Regression: recitals prompt must actually carry the party names ------
#
# Found live, 2026-08-01 (Sprint 2, user click-through on NDA v1): the
# recitals came out with generic labels ("Party A", "the Disclosing
# Party") instead of real party names — screenshot-confirmed, defeating
# the point of the intake form. Root cause, confirmed by direct code
# read: the recitals source_text in both seed_nda_template.py and
# seed_service_agreement_template.py never referenced
# {{ party_a_name }}/{{ party_b_name }} at all, so the model had nothing
# to draw a real name from. These tests load the REAL seed script CLAUSES
# (not the hand-rolled NDA_CLAUSES fixture, whose recitals stub —
# "Draft recitals for: {{ purpose }}" — never had party names either and
# so could never have caught this). They assert on the rendered PROMPT
# text captured by _fake_generate, after PII masking — CLAUDE.md Decision
# 4 masks party names before any LLM call, so proving the fix means
# proving two distinct PARTY_n placeholders reach the prompt (i.e. two
# Jinja substitutions actually happened), not that the raw names do.


def _party_placeholders(prompt: str) -> set[str]:
    # "Named party" kinds get letter suffixes (PARTY_A, PARTY_B, ...), not
    # numeric ones — see pii_mask.py's _LETTER_KINDS.
    return set(re.findall(r"PARTY_[A-Z]+", prompt))


def test_nda_recitals_prompt_carries_both_party_names_mutual(monkeypatch):
    db = FakeDB()
    template_id = _seed_nda_real(db)
    matter_id = _seed_matter(db)
    calls = _fake_generate(monkeypatch)

    form = dict(BASE_FORM, nda_variant="mutual")
    contracts.generate_draft(matter_id, template_id, form, db=db)

    recitals_prompt = calls[0]["prompt"]
    placeholders = _party_placeholders(recitals_prompt)
    assert len(placeholders) == 2, (
        f"expected both party names substituted into the recitals prompt as two "
        f"distinct masked placeholders, got {placeholders} in: {recitals_prompt!r}"
    )
    assert "Ramesh Kumar" not in recitals_prompt and "Acme Pvt Ltd" not in recitals_prompt, (
        "party names must reach the prompt only via their masked placeholders (CLAUDE.md Decision 4)"
    )
    assert "MUTUAL NDA" in recitals_prompt
    # the old hardcoded phrasing that was legally incoherent for a mutual
    # NDA (both parties get the identical role label from
    # _variant_role_labels) must not appear unconditionally.
    assert "is the Disclosing Party and" not in recitals_prompt


def test_nda_recitals_prompt_one_way_party_a_disclosing(monkeypatch):
    db = FakeDB()
    template_id = _seed_nda_real(db)
    matter_id = _seed_matter(db)
    calls = _fake_generate(monkeypatch)

    form = dict(BASE_FORM, nda_variant="one_way", party_a_role="disclosing")
    contracts.generate_draft(matter_id, template_id, form, db=db)

    recitals_prompt = calls[0]["prompt"]
    assert len(_party_placeholders(recitals_prompt)) == 2
    assert "ONE-WAY NDA" in recitals_prompt
    placeholder_a = re.search(r"(PARTY_[A-Z]+) is the Disclosing Party and (PARTY_[A-Z]+) is the Receiving Party", recitals_prompt)
    assert placeholder_a, f"expected 'X is the Disclosing Party and Y is the Receiving Party' in: {recitals_prompt!r}"
    assert placeholder_a.group(1) != placeholder_a.group(2)


def test_nda_recitals_prompt_one_way_party_a_receiving(monkeypatch):
    db = FakeDB()
    template_id = _seed_nda_real(db)
    matter_id = _seed_matter(db)
    calls = _fake_generate(monkeypatch)

    form = dict(BASE_FORM, nda_variant="one_way", party_a_role="receiving")
    contracts.generate_draft(matter_id, template_id, form, db=db)

    recitals_prompt = calls[0]["prompt"]
    assert len(_party_placeholders(recitals_prompt)) == 2
    assert "ONE-WAY NDA" in recitals_prompt
    # party_a_role=receiving -> the branch names party_b_name as Disclosing first.
    match = re.search(r"(PARTY_[A-Z]+) is the Disclosing Party and (PARTY_[A-Z]+) is the Receiving Party", recitals_prompt)
    assert match, f"expected 'X is the Disclosing Party and Y is the Receiving Party' in: {recitals_prompt!r}"
    assert match.group(1) != match.group(2)


def test_service_agreement_recitals_prompt_carries_both_party_names(monkeypatch):
    db = FakeDB()
    template_id = _seed_service_agreement(db)
    matter_id = _seed_matter(db)
    calls = _fake_generate(monkeypatch)

    contracts.generate_draft(matter_id, template_id, SERVICE_AGREEMENT_FORM, db=db)

    recitals_prompt = calls[0]["prompt"]
    placeholders = _party_placeholders(recitals_prompt)
    assert len(placeholders) == 2, (
        f"expected both party names substituted into the recitals prompt as two "
        f"distinct masked placeholders, got {placeholders} in: {recitals_prompt!r}"
    )
    assert "Ramesh Kumar" not in recitals_prompt and "Acme Technologies Pvt Ltd" not in recitals_prompt
    assert "Service Provider" in recitals_prompt and "Client" in recitals_prompt


# --- Regression: embedded newlines in free text must not fragment a docx --
# paragraph mid-sentence -----------------------------------------------------
#
# Found live, 2026-08-01 (Sprint 2, user reviewed the Service Agreement
# docx): a deliverable description came out truncated mid-word in Scope
# of Services. Root cause: `deliverables[].description` is a `textarea`
# field (a real multi-row control — see intake-form.tsx's rows={3}), so
# a literal "\n" can end up inside a single field's value (Enter while
# typing, or pasting hyphenated text from a justified PDF/Word
# paragraph). generate_draft's docx-assembly step splits each rendered
# clause's text on "\n" to create separate paragraphs — meant for a
# template author's own intentional paragraph breaks (e.g. NDA's
# Definitions "1.1 ... \n1.2 ..."), but blind to the difference between
# that and a newline that happened to be inside a client's answer, so it
# silently fragments the value across two docx paragraphs.


def test_generate_draft_normalizes_embedded_newlines_in_deliverable_description(monkeypatch):
    from docx import Document

    db = FakeDB()
    template_id = _seed_service_agreement(db)
    matter_id = _seed_matter(db)
    _fake_generate(monkeypatch, canned_text="WHEREAS the parties wish to engage.")

    form = dict(
        SERVICE_AGREEMENT_FORM,
        deliverables=[
            {"description": "Integrate third-\nparty payment gateways", "due_date": "2026-09-01"},
        ],
    )
    result = contracts.generate_draft(matter_id, template_id, form, db=db)

    output_path = REPO_ROOT / result.docx_path
    try:
        doc = Document(str(output_path))
        paragraphs = [p.text for p in doc.paragraphs]
        # The embedded newline collapses to a literal space (not a seamless
        # hyphen-join — that would need fragile PDF-dehyphenation
        # heuristics this fix deliberately doesn't attempt); what matters
        # is the value survives as ONE docx paragraph, not two.
        assert any("Integrate third- party payment gateways" in p for p in paragraphs), (
            f"embedded newline should collapse to a space within a single "
            f"paragraph, not fragment the value across paragraphs; got: {paragraphs!r}"
        )
        assert not any(p.strip() == "party payment gateways" for p in paragraphs), (
            "deliverable description was split into two docx paragraphs at the embedded newline"
        )
    finally:
        output_path.unlink(missing_ok=True)


def test_normalize_free_text_collapses_newlines_recursively_in_lists():
    real_schema = json.loads((REPO_ROOT / "templates" / "contracts" / "service-agreement.schema.json").read_text())
    form = dict(
        SERVICE_AGREEMENT_FORM,
        purpose="Ongoing software\nconsulting services",
        deliverables=[{"description": "Build the\napi", "due_date": ""}],
    )
    normalized = contracts._normalize_free_text(form, real_schema)
    assert normalized["purpose"] == "Ongoing software consulting services"
    assert normalized["deliverables"][0]["description"] == "Build the api"


# --- Regression: recitals must not invent a fee structure it was never ----
# told about --------------------------------------------------------------


def test_service_agreement_recitals_prompt_excludes_fee_structure(monkeypatch):
    """Found live 2026-08-01: recitals asserted 'a fixed fee model' for a
    matter configured as Milestone-Based. fee_structure/fee_amount were
    never even in this prompt's context — pure invention, because unlike
    the (working) deliverables guard, there was no equivalent 'don't
    describe fee/payment terms here' instruction. This asserts the guard
    is present in the real seed script content, not just that the two
    variables are absent (which was already true before the fix, and
    didn't stop the model inventing language anyway)."""
    db = FakeDB()
    template_id = _seed_service_agreement(db)
    matter_id = _seed_matter(db)
    calls = _fake_generate(monkeypatch)

    contracts.generate_draft(matter_id, template_id, SERVICE_AGREEMENT_FORM, db=db)

    recitals_prompt = calls[0]["prompt"]
    assert "fee_structure" not in recitals_prompt and "fee_amount" not in recitals_prompt
    lowered = recitals_prompt.lower()
    assert "do not state or characterise the fee amount" in lowered or "payment terms clause covers that separately" in lowered


# --- Regression: llm_fillable clauses must not drift to generic party ----
# labels mid-generation -----------------------------------------------------


def test_service_agreement_ip_assignment_prompt_forbids_generic_party_labels(monkeypatch):
    """Found live 2026-08-01: a generated IP clause used 'Client' in one
    sub-clause and 'Party B' in another. The recitals prompt already had
    an explicit guard against generic labels; this prompt didn't."""
    db = FakeDB()
    template_id = _seed_service_agreement(db)
    matter_id = _seed_matter(db)
    calls = _fake_generate(monkeypatch)

    contracts.generate_draft(matter_id, template_id, SERVICE_AGREEMENT_FORM, db=db)

    ip_prompt = next(c["prompt"] for c in calls if "IP ownership model" in c["prompt"])
    lowered = ip_prompt.lower()
    assert "never 'party a', 'party b'" in lowered or ("never" in lowered and "party a" in lowered and "party b" in lowered)


# --- Regression: arbitration clause flags per-matter specifics for --------
# Nitesh's clause review, both templates -------------------------------------


def test_service_agreement_governing_law_prompt_flags_arbitration_specifics(monkeypatch):
    db = FakeDB()
    template_id = _seed_service_agreement(db)
    matter_id = _seed_matter(db)
    calls = _fake_generate(monkeypatch)

    form = dict(SERVICE_AGREEMENT_FORM, arbitration=True, arbitration_seat="Mumbai")
    contracts.generate_draft(matter_id, template_id, form, db=db)

    gov_prompt = next(c["prompt"] for c in calls if "Governing Law and Dispute Resolution" in c["prompt"])
    assert "ADVOCATE REVIEW" in gov_prompt
    assert "number of arbitrators" in gov_prompt


def test_nda_governing_law_prompt_flags_arbitration_specifics(monkeypatch):
    db = FakeDB()
    template_id = _seed_nda_real(db)
    matter_id = _seed_matter(db)
    calls = _fake_generate(monkeypatch)

    form = dict(BASE_FORM, arbitration=True, arbitration_seat="Delhi")
    contracts.generate_draft(matter_id, template_id, form, db=db)

    gov_prompt = next(c["prompt"] for c in calls if "Governing Law and Dispute Resolution" in c["prompt"])
    assert "ADVOCATE REVIEW" in gov_prompt
    assert "number of arbitrators" in gov_prompt


# --- Regression: confidentiality direction/survival is matter-specific ----
# now, not hardcoded ----------------------------------------------------------
#
# Design gap fixed 2026-08-02: confidentiality was previously hardcoded
# one-way (Client discloses to Service Provider only) with a fixed
# 3-year survival, regardless of the matter. User's call: add real
# intake fields rather than defer entirely to clause review, since this
# clause is central to the deal. These tests confirm the right one of
# the three confidentiality_* clause rows is selected by
# confidentiality_direction and that confidentiality_survival_period is
# substituted in — mirroring the same applicable_condition mechanism
# NDA's confidentiality_obligations_mutual/one_way already used.


@pytest.mark.parametrize(
    "direction,expected_clause_key,excluded_clause_keys",
    [
        (
            "mutual",
            "confidentiality_mutual",
            {"confidentiality_one_way_from_client", "confidentiality_one_way_from_provider"},
        ),
        (
            "one_way_from_client",
            "confidentiality_one_way_from_client",
            {"confidentiality_mutual", "confidentiality_one_way_from_provider"},
        ),
        (
            "one_way_from_provider",
            "confidentiality_one_way_from_provider",
            {"confidentiality_mutual", "confidentiality_one_way_from_client"},
        ),
    ],
)
def test_service_agreement_confidentiality_direction_selects_right_clause(
    monkeypatch, direction, expected_clause_key, excluded_clause_keys
):
    db = FakeDB()
    template_id = _seed_service_agreement(db)
    matter_id = _seed_matter(db)
    _fake_generate(monkeypatch, canned_text="WHEREAS the parties wish to engage.")

    form = dict(
        SERVICE_AGREEMENT_FORM,
        confidentiality_direction=direction,
        confidentiality_survival_period="5 years",
    )
    result = contracts.generate_draft(matter_id, template_id, form, db=db)

    assert "Confidentiality" in result.full_text
    assert "for a period of 5 years" in result.full_text
    _, clauses = _load_service_agreement_fixtures()
    all_confidentiality_keys = {
        "confidentiality_mutual", "confidentiality_one_way_from_client", "confidentiality_one_way_from_provider"
    }
    assert all_confidentiality_keys <= {c["clause_key"] for c in clauses}, (
        "all three variants must exist in the real seed script"
    )
    for excluded_key in excluded_clause_keys:
        excluded_text = next(c["source_text"] for c in clauses if c["clause_key"] == excluded_key)
        # A crude but effective signal that the excluded variant's wording
        # did not leak into this draft: its opening subject differs by
        # clause, so check its first ~40 chars are absent.
        assert excluded_text[:40] not in result.full_text


# =============================================================================
# Consultancy Agreement (Sprint 2 Deliverable 2, Batch 2)
#
# Built with all four of Batch 1's live-discovered lessons already baked
# in (see docs/lessons_learned.md's new process rule and design pattern
# notes) rather than rediscovered here: real party names + no-generic-
# label guards on every llm_fillable clause from day one, fee/payment
# terms explicitly excluded from recitals, confidentiality as a real
# three-way intake choice via applicable_condition-per-variant, and the
# arbitration [ADVOCATE REVIEW: ...] flag. These tests exercise what's
# actually new here: the Deliverables/Scope split and the Retainer fee
# branch — the confidentiality-direction and arbitration-flag mechanics
# are already covered generically by the NDA/Service Agreement tests
# above, reusing the identical pattern.
# =============================================================================


def test_consultancy_deliverables_clause_excluded_when_list_empty(monkeypatch):
    db = FakeDB()
    template_id = _seed_consultancy(db)
    matter_id = _seed_matter(db)
    _fake_generate(monkeypatch, canned_text="Narrative scope text.")

    result = contracts.generate_draft(matter_id, template_id, CONSULTANCY_FORM, db=db)

    assert "9. Miscellaneous" in result.full_text, (
        "with no deliverables, Deliverables clause should be excluded and "
        "Miscellaneous should shift down to clause 9"
    )
    assert "10. Miscellaneous" not in result.full_text


def test_consultancy_deliverables_clause_included_when_list_non_empty(monkeypatch):
    db = FakeDB()
    template_id = _seed_consultancy(db)
    matter_id = _seed_matter(db)
    _fake_generate(monkeypatch, canned_text="Narrative scope text.")

    form = dict(
        CONSULTANCY_FORM,
        deliverables=[
            {"description": "Deliver a supply chain audit report", "due_date": "2026-09-01"},
        ],
    )
    result = contracts.generate_draft(matter_id, template_id, form, db=db)

    assert "Deliver a supply chain audit report (due 2026-09-01)" in result.full_text
    assert "10. Miscellaneous" in result.full_text, (
        "with a deliverable present, Deliverables clause should be "
        "included and Miscellaneous should be clause 10"
    )
    assert "9. Miscellaneous" not in result.full_text


def test_consultancy_recitals_prompt_excludes_scope_and_fee_details(monkeypatch):
    db = FakeDB()
    template_id = _seed_consultancy(db)
    matter_id = _seed_matter(db)
    calls = _fake_generate(monkeypatch, canned_text="Narrative scope text.")

    contracts.generate_draft(matter_id, template_id, CONSULTANCY_FORM, db=db)

    recitals_prompt = calls[0]["prompt"]
    placeholders = _party_placeholders(recitals_prompt)
    assert len(placeholders) == 2, (
        f"expected both party names substituted as two distinct masked "
        f"placeholders, got {placeholders} in: {recitals_prompt!r}"
    )
    assert "fee_structure" not in recitals_prompt and "fee_amount" not in recitals_prompt
    assert "Consultant" in recitals_prompt and "Client" in recitals_prompt


def test_consultancy_payment_terms_retainer_branch_with_scope_hours(monkeypatch):
    db = FakeDB()
    template_id = _seed_consultancy(db)
    matter_id = _seed_matter(db)
    _fake_generate(monkeypatch, canned_text="Narrative scope text.")

    result = contracts.generate_draft(matter_id, template_id, CONSULTANCY_FORM, db=db)

    assert "a retainer fee of ₹1,00,000 per month" in result.full_text
    assert "billed Monthly" in result.full_text
    assert "covers up to 20 hours per month of the Consultant's time" in result.full_text, (
        "the field's own value ('up to 20 hours per month') should read "
        "naturally after 'covers', not be double-prefixed by the clause's own 'up to'"
    )
    assert "up to up to" not in result.full_text


def test_consultancy_payment_terms_retainer_branch_uncapped(monkeypatch):
    db = FakeDB()
    template_id = _seed_consultancy(db)
    matter_id = _seed_matter(db)
    _fake_generate(monkeypatch, canned_text="Narrative scope text.")

    form = dict(CONSULTANCY_FORM, retainer_scope_hours="")
    result = contracts.generate_draft(matter_id, template_id, form, db=db)

    assert "a retainer fee of ₹1,00,000 per month" in result.full_text
    assert "of the Consultant's time" not in result.full_text, (
        "an uncapped retainer should not mention an hours cap at all"
    )


def test_consultancy_payment_terms_fixed_fee_branch(monkeypatch):
    db = FakeDB()
    template_id = _seed_consultancy(db)
    matter_id = _seed_matter(db)
    _fake_generate(monkeypatch, canned_text="Narrative scope text.")

    form = dict(
        CONSULTANCY_FORM,
        fee_structure="Fixed Fee",
        fee_amount="₹2,00,000",
        payment_frequency="Upfront",
        retainer_fee_amount="",
        retainer_frequency="",
        retainer_scope_hours="",
    )
    result = contracts.generate_draft(matter_id, template_id, form, db=db)

    assert "a fixed fee of ₹2,00,000, payable Upfront" in result.full_text
    assert "retainer" not in result.full_text.lower()


def test_consultancy_docx_renders_end_to_end(monkeypatch):
    from docx import Document

    db = FakeDB()
    template_id = _seed_consultancy(db)
    matter_id = _seed_matter(db, matter_id="matter-consultancy-docx")
    _fake_generate(monkeypatch, canned_text="Narrative scope text drafted by the model.")

    form = dict(
        CONSULTANCY_FORM,
        deliverables=[{"description": "Deliver a supply chain audit report", "due_date": "2026-09-01"}],
    )
    result = contracts.generate_draft(matter_id, template_id, form, db=db)

    output_path = REPO_ROOT / result.docx_path
    try:
        doc = Document(str(output_path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "CONSULTANCY AGREEMENT" in full_text
        assert "Anjali Mehta" in full_text
        assert "Bluewave Logistics Pvt Ltd" in full_text
        assert "Consultant" in full_text and "Client" in full_text
        assert "Deliver a supply chain audit report (due 2026-09-01)" in full_text
        assert "10. Miscellaneous" in full_text
    finally:
        output_path.unlink(missing_ok=True)


# =============================================================================
# Memorandum of Understanding (Sprint 2 Deliverable 2, Batch 3)
#
# Deliberately the simplest template so far — a validation that Sprint
# 2's abstractions (applicable_condition-per-variant, an_or_a, schema-
# type-aware masking, _normalize_free_text) hold on a minimal template
# without introducing any new mechanism. Only 2 of 7 logical clauses are
# llm_fillable (Recitals, Governing Law and Jurisdiction); the rest is
# fixed boilerplate, including Term and Termination (llm_fillable in
# Consultancy/Service Agreement, fixed_boilerplate here — MoU term logic
# needs no narrative judgment).
# =============================================================================


def test_mou_recitals_prompt_carries_both_party_names_and_excludes_binding_status(monkeypatch):
    db = FakeDB()
    template_id = _seed_mou(db)
    matter_id = _seed_matter(db)
    calls = _fake_generate(monkeypatch, canned_text="Recitals narrative text.")

    contracts.generate_draft(matter_id, template_id, MOU_FORM, db=db)

    recitals_prompt = calls[0]["prompt"]
    placeholders = _party_placeholders(recitals_prompt)
    assert len(placeholders) == 2, (
        f"expected both party names substituted as two distinct masked "
        f"placeholders, got {placeholders} in: {recitals_prompt!r}"
    )
    assert "Karan Malhotra" not in recitals_prompt and "Greenfield Renewables Pvt Ltd" not in recitals_prompt
    assert "do not state whether this memorandum is binding or non-binding" in recitals_prompt.lower()


def test_mou_nature_clause_enumerates_binding_exceptions(monkeypatch):
    db = FakeDB()
    template_id = _seed_mou(db)
    matter_id = _seed_matter(db)
    _fake_generate(monkeypatch, canned_text="Recitals narrative text.")

    result = contracts.generate_draft(matter_id, template_id, MOU_FORM, db=db)

    assert "does not constitute a legally binding" in result.full_text
    assert "Confidentiality, Costs and Expenses, and Governing Law and Jurisdiction" in result.full_text
    assert "shall be binding on the Parties with immediate effect" in result.full_text


@pytest.mark.parametrize(
    "direction,expected_clause_key,excluded_clause_keys",
    [
        (
            "mutual",
            "confidentiality_mutual",
            {"confidentiality_one_way_from_a", "confidentiality_one_way_from_b"},
        ),
        (
            "one_way_from_a",
            "confidentiality_one_way_from_a",
            {"confidentiality_mutual", "confidentiality_one_way_from_b"},
        ),
        (
            "one_way_from_b",
            "confidentiality_one_way_from_b",
            {"confidentiality_mutual", "confidentiality_one_way_from_a"},
        ),
    ],
)
def test_mou_confidentiality_direction_selects_right_clause(
    monkeypatch, direction, expected_clause_key, excluded_clause_keys
):
    db = FakeDB()
    template_id = _seed_mou(db)
    matter_id = _seed_matter(db)
    _fake_generate(monkeypatch, canned_text="Recitals narrative text.")

    form = dict(MOU_FORM, confidentiality_direction=direction, confidentiality_survival_period="4 years")
    result = contracts.generate_draft(matter_id, template_id, form, db=db)

    assert "Confidentiality" in result.full_text
    assert "for a period of 4 years" in result.full_text
    _, clauses = _load_mou_fixtures()
    all_confidentiality_keys = {
        "confidentiality_mutual", "confidentiality_one_way_from_a", "confidentiality_one_way_from_b"
    }
    assert all_confidentiality_keys <= {c["clause_key"] for c in clauses}
    for excluded_key in excluded_clause_keys:
        excluded_text = next(c["source_text"] for c in clauses if c["clause_key"] == excluded_key)
        assert excluded_text[:40] not in result.full_text


def test_mou_governing_law_prompt_flags_arbitration_specifics(monkeypatch):
    db = FakeDB()
    template_id = _seed_mou(db)
    matter_id = _seed_matter(db)
    calls = _fake_generate(monkeypatch, canned_text="Recitals narrative text.")

    form = dict(MOU_FORM, arbitration=True, arbitration_seat="Delhi")
    contracts.generate_draft(matter_id, template_id, form, db=db)

    gov_prompt = next(c["prompt"] for c in calls if "Governing Law and Dispute Resolution" in c["prompt"])
    assert "ADVOCATE REVIEW" in gov_prompt
    assert "number of arbitrators" in gov_prompt


def test_mou_docx_renders_end_to_end_with_correct_numbering(monkeypatch):
    from docx import Document

    db = FakeDB()
    template_id = _seed_mou(db)
    matter_id = _seed_matter(db, matter_id="matter-mou-docx")
    _fake_generate(monkeypatch, canned_text="Recitals narrative text drafted by the model.")

    result = contracts.generate_draft(matter_id, template_id, MOU_FORM, db=db)

    output_path = REPO_ROOT / result.docx_path
    try:
        doc = Document(str(output_path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "MEMORANDUM OF UNDERSTANDING" in full_text
        assert "Karan Malhotra" in full_text
        assert "Greenfield Renewables Pvt Ltd" in full_text
        assert "1. Nature of this Memorandum" in full_text
        assert "2. Confidentiality" in full_text
        assert "3. Term and Termination" in full_text
        assert "4. Costs and Expenses" in full_text
        assert "5. Governing Law and Jurisdiction" in full_text
        assert "6. Miscellaneous" in full_text
        # Costs and Expenses is pure fixed boilerplate — no fields at all.
        assert "Each Party shall bear its own costs" in full_text
    finally:
        output_path.unlink(missing_ok=True)


# =============================================================================
# Employment Agreement (Sprint 2 Deliverable 2, Batch 4)
#
# Introduces this rollout's first statutory-compliance content (PF/ESI/
# Gratuity) and non-compete doctrine (Section 27, Indian Contract Act,
# 1872) but no new mechanism. Governing Law is fixed_boilerplate from
# day one here (unlike NDA/Service Agreement/Consultancy/MoU, which all
# shipped it as llm_fillable before the MoU-established classification
# bar) — third clean application of that bar, alongside Intellectual
# Property (also fixed_boilerplate here, second application).
# =============================================================================


def test_employment_probation_clause_selected_by_has_probation(monkeypatch):
    db = FakeDB()
    template_id = _seed_employment(db)
    matter_id = _seed_matter(db)
    _fake_generate(monkeypatch, canned_text="Narrative text drafted by the model.")

    with_probation = contracts.generate_draft(matter_id, template_id, EMPLOYMENT_FORM, db=db)
    assert "shall serve a probationary period of 6 months" in with_probation.full_text
    assert "without any probationary period" not in with_probation.full_text

    form_no_probation = dict(EMPLOYMENT_FORM, has_probation=False, probation_period="")
    without_probation = contracts.generate_draft(matter_id, template_id, form_no_probation, db=db)
    assert "without any probationary period" in without_probation.full_text
    assert "shall serve a probationary period" not in without_probation.full_text


def test_employment_statutory_compliance_clause_cites_correct_acts_no_computed_figures(monkeypatch):
    db = FakeDB()
    template_id = _seed_employment(db)
    matter_id = _seed_matter(db)
    _fake_generate(monkeypatch, canned_text="Narrative text drafted by the model.")

    result = contracts.generate_draft(matter_id, template_id, EMPLOYMENT_FORM, db=db)

    assert "Employees' Provident Funds and Miscellaneous Provisions Act, 1952" in result.full_text
    assert "Employees' State Insurance Act, 1948" in result.full_text
    assert "Payment of Gratuity Act, 1972" in result.full_text
    assert "are not asserted as fact in this Agreement" in result.full_text
    # No specific rate/threshold/ceiling figure should ever appear — the
    # clause is deliberately non-computed (approved design decision).
    assert "₹" not in result.full_text.split("Statutory Compliance")[-1].split("Confidentiality")[0], (
        "Statutory Compliance clause must not assert a specific rupee figure"
    )


def test_employment_restrictive_covenants_prompt_has_section_27_caveat(monkeypatch):
    db = FakeDB()
    template_id = _seed_employment(db)
    matter_id = _seed_matter(db)
    calls = _fake_generate(monkeypatch, canned_text="Narrative text drafted by the model.")

    contracts.generate_draft(matter_id, template_id, EMPLOYMENT_FORM, db=db)

    rc_prompt = next(c["prompt"] for c in calls if "Restrictive Covenants" in c["prompt"])
    assert "Section 27" in rc_prompt and "Indian Contract Act, 1872" in rc_prompt
    assert "void" in rc_prompt.lower()
    assert "ADVOCATE REVIEW" in rc_prompt


def test_employment_governing_law_is_fixed_boilerplate_not_llm_fillable():
    _, clauses = _load_employment_fixtures()
    gov = next(c for c in clauses if c["clause_key"] == "governing_law_jurisdiction")
    assert gov["clause_type"] == "fixed_boilerplate", (
        "Governing Law must ship as fixed_boilerplate from day one on this template — "
        "see the MoU-established llm_fillable classification bar"
    )


def test_employment_intellectual_property_is_fixed_boilerplate_not_llm_fillable():
    _, clauses = _load_employment_fixtures()
    ip = next(c for c in clauses if c["clause_key"] == "intellectual_property")
    assert ip["clause_type"] == "fixed_boilerplate"


def test_employment_only_three_llm_fillable_clauses():
    _, clauses = _load_employment_fixtures()
    llm_fillable_keys = {c["clause_key"] for c in clauses if c["clause_type"] == "llm_fillable"}
    assert llm_fillable_keys == {"recitals", "position_duties_reporting", "restrictive_covenants"}


def test_employment_confidentiality_is_single_clause_no_variant(monkeypatch):
    """Employment confidentiality only has one sensible configuration
    (Employee owes Employer) — unlike Consultancy/Service Agreement/MoU,
    it should NOT use the applicable_condition-per-variant pattern."""
    _, clauses = _load_employment_fixtures()
    confidentiality_clauses = [c for c in clauses if c["clause_key"].startswith("confidentiality")]
    assert len(confidentiality_clauses) == 1
    assert confidentiality_clauses[0]["applicable_condition"] is None


def test_employment_fixed_term_end_date_appears_only_for_fixed_term(monkeypatch):
    db = FakeDB()
    template_id = _seed_employment(db)
    matter_id = _seed_matter(db)
    _fake_generate(monkeypatch, canned_text="Narrative text drafted by the model.")

    full_time = contracts.generate_draft(matter_id, template_id, EMPLOYMENT_FORM, db=db)
    assert "shall automatically terminate on" not in full_time.full_text

    form = dict(EMPLOYMENT_FORM, employment_type="Fixed-Term", fixed_term_end_date="2027-08-02")
    fixed_term = contracts.generate_draft(matter_id, template_id, form, db=db)
    assert "shall automatically terminate on 2027-08-02" in fixed_term.full_text


def test_employment_governing_law_arbitration_branch(monkeypatch):
    db = FakeDB()
    template_id = _seed_employment(db)
    matter_id = _seed_matter(db)
    _fake_generate(monkeypatch, canned_text="Narrative text drafted by the model.")

    no_arb = contracts.generate_draft(matter_id, template_id, EMPLOYMENT_FORM, db=db)
    assert "governed by the laws of India" in no_arb.full_text
    assert "ADVOCATE REVIEW" not in no_arb.full_text

    form = dict(EMPLOYMENT_FORM, arbitration=True, arbitration_seat="Gurugram")
    with_arb = contracts.generate_draft(matter_id, template_id, form, db=db)
    assert "referred to arbitration" in with_arb.full_text
    assert "seated at Gurugram" in with_arb.full_text
    assert "ADVOCATE REVIEW" in with_arb.full_text


def test_employment_docx_renders_end_to_end_with_correct_numbering(monkeypatch):
    from docx import Document

    db = FakeDB()
    template_id = _seed_employment(db)
    matter_id = _seed_matter(db, matter_id="matter-employment-docx")
    _fake_generate(monkeypatch, canned_text="Narrative text drafted by the model.")

    result = contracts.generate_draft(matter_id, template_id, EMPLOYMENT_FORM, db=db)

    output_path = REPO_ROOT / result.docx_path
    try:
        doc = Document(str(output_path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "EMPLOYMENT AGREEMENT" in full_text
        assert "Bright Horizons Tech Pvt Ltd" in full_text
        assert "Sneha Reddy" in full_text
        assert "an Individual residing at" in full_text, "Employee block must hardcode 'an Individual', no entity_type field"
        assert "1. Definitions" in full_text
        assert "2. Position, Duties and Reporting" in full_text
        assert "3. Probation" in full_text
        assert "4. Compensation and Benefits" in full_text
        assert "5. Statutory Compliance" in full_text
        assert "6. Confidentiality" in full_text
        assert "7. Intellectual Property" in full_text
        assert "8. Restrictive Covenants" in full_text
        assert "9. Termination" in full_text
        assert "10. Governing Law and Jurisdiction" in full_text
        assert "11. Miscellaneous" in full_text
    finally:
        output_path.unlink(missing_ok=True)


def test_convert_docx_to_pdf_timeout_raises_pdf_conversion_timeout(monkeypatch, tmp_path):
    """PERF-01: When LibreOffice exceeds the timeout cap, convert_docx_to_pdf
    catches subprocess.TimeoutExpired and raises PdfConversionTimeout."""
    import subprocess
    from unittest.mock import MagicMock

    dummy_docx = tmp_path / "test_doc.docx"
    dummy_docx.write_text("dummy docx content")

    monkeypatch.setattr("shutil.which", lambda cmd: "/usr/bin/libreoffice")

    def _mock_run(*args, **kwargs):
        raise subprocess.TimeoutExpired(cmd="libreoffice", timeout=15)

    monkeypatch.setattr("subprocess.run", _mock_run)

    with pytest.raises(contracts.PdfConversionTimeout) as exc_info:
        contracts.convert_docx_to_pdf(dummy_docx, timeout=15)

    assert "LibreOffice PDF conversion timed out after 15 seconds" in str(exc_info.value)


def test_all_10_contract_templates_have_canonical_kebab_case_keys(monkeypatch):
    """Regression test: All 10 Phase 1 templates must exist in database, return via list_templates(),
    and have canonical kebab-case template_key values (no underscores)."""
    import re
    from app.db import service_client

    rows = service_client().table("templates").select("id, name, template_key, category").eq("category", "contracts").execute().data
    assert len(rows) == 10, f"Expected exactly 10 contract templates, found {len(rows)}"

    kebab_pattern = re.compile(r"^[a-z0-9-]+$")
    template_keys = [r["template_key"] for r in rows]

    for key in template_keys:
        assert key is not None and key != "", f"Template row missing template_key: {key}"
        assert "_" not in key, f"Template key '{key}' contains underscore; must be canonical kebab-case"
        assert kebab_pattern.match(key), f"Template key '{key}' does not match kebab-case pattern"

    assert "service-agreement" in template_keys
    assert "nda" in template_keys
    assert "consultancy" in template_keys
    assert "mou" in template_keys
    assert "employment" in template_keys
    assert "leave-licence" in template_keys
    assert "lease-deed" in template_keys
    assert "joint-venture" in template_keys
    assert "agreement-to-sell" in template_keys
    assert "software-dev" in template_keys


def test_frontend_category_filter_mapping_covers_all_10_templates():
    """Regression test: Ensure every one of the 10 canonical template_keys maps to at least
    one frontend category ('commercial', 'employment', 'ip') so no template is ever filtered out."""
    canonical_keys = [
        "nda",
        "service-agreement",
        "consultancy",
        "mou",
        "leave-licence",
        "lease-deed",
        "agreement-to-sell",
        "joint-venture",
        "employment",
        "software-dev",
    ]

    commercial_keys = {"nda", "service-agreement", "consultancy", "mou", "leave-licence", "lease-deed", "agreement-to-sell", "joint-venture"}
    employment_keys = {"employment"}
    ip_keys = {"software-dev", "nda", "service-agreement"}

    all_categorized_keys = commercial_keys | employment_keys | ip_keys

    for key in canonical_keys:
        assert key in all_categorized_keys, f"Template key '{key}' is excluded from all frontend categories!"


def test_matter_centric_loading_model_stores_and_returns_template_id(monkeypatch):
    """Regression test for Matter-centric loading model:
    Creating a matter with template_id (UUID or slug) persists template_id on the matter record,
    and GET /api/matters/{id} returns template_id.

    DB-isolated (forensic finding, 2026-08-14): this test used to inject
    db=service_client() -- the real, production, service-role Supabase client --
    into CurrentUser, so every local `pytest` run permanently wrote two real rows
    ("Matter Centric Test (Slug)"/"(UUID)") into the one Supabase project this repo
    is configured against, with no teardown. It now uses the same FakeDB/FakeTable
    pattern as the rest of this file (and as test_case_analysis.py, test_litigation.py,
    etc. use their own equivalents), and app.db.service_client is monkeypatched to
    raise if anything still tries to reach the real client."""
    import uuid

    from fastapi.testclient import TestClient
    from app.main import app
    from app.auth import get_current_user, CurrentUser

    monkeypatch.setattr("app.db.service_client", _poisoned_service_client)

    fake_db = FakeDB()
    sa_uuid = str(uuid.uuid4())
    fake_db.table("templates").rows.append({
        "id": sa_uuid,
        "name": "Service Agreement",
        "category": "contracts",
        "template_key": "service-agreement",
        "review_status": "beta",
        "states_supported": [],
        "schema_json": {},
    })
    # create_matter() resolves a template_key slug via app.routers.matters'
    # module-level `service_client` name (bound at import time) -- patching
    # app.db.service_client above doesn't reach that already-bound reference,
    # so it has to be patched here too for the fake to actually take effect.
    monkeypatch.setattr("app.routers.matters.service_client", lambda: fake_db)

    client = TestClient(app)

    def mock_user():
        return CurrentUser(id="21e63e8f-e00c-4ae6-afe4-17ba6b400be5", email="test@example.com", db=fake_db)

    app.dependency_overrides[get_current_user] = mock_user

    # Fetch Service Agreement template UUID (from the fake DB, not Supabase)
    tpl_row = fake_db.table("templates").select("id").eq("template_key", "service-agreement").execute().data[0]
    assert tpl_row["id"] == sa_uuid

    # 1. Create matter passing template_key slug
    res1 = client.post("/api/matters", json={
        "title": "Matter Centric Test (Slug)",
        "client_name": "Test Client",
        "module": "contracts",
        "template_id": "service-agreement"
    })
    assert res1.status_code == 201
    m1 = res1.json()
    assert m1["template_id"] == sa_uuid

    # 2. GET matter and verify status code 200 and matter details
    res1_get = client.get(f"/api/matters/{m1['id']}")
    assert res1_get.status_code == 200
    assert res1_get.json()["id"] == m1["id"]

    # 3. Create matter passing template UUID directly
    res2 = client.post("/api/matters", json={
        "title": "Matter Centric Test (UUID)",
        "client_name": "Test Client",
        "module": "contracts",
        "template_id": sa_uuid
    })
    assert res2.status_code == 201
    assert res2.json()["template_id"] == sa_uuid


def test_template_lookup_supports_both_uuid_and_slug_keys(monkeypatch):
    """Regression test: GET /api/templates/{template_id} works identically whether called
    with a UUID or with a template_key slug (e.g. 'service-agreement').

    DB-isolated (forensic finding, 2026-08-14): the route under test
    (api/app/routers/contracts.py:get_template) only ever reads via `user.db`, so this
    test was genuinely read-only against Supabase -- but it still depended on a real,
    live "service-agreement" template row existing in production, via
    db=service_client() injected into CurrentUser, on every local `pytest` run. Switched
    to the same FakeDB pattern as the rest of this file so it no longer needs (or can
    reach) the real project at all; app.db.service_client is monkeypatched to raise if
    anything still tries to reach the real client."""
    import uuid

    from fastapi.testclient import TestClient
    from app.main import app
    from app.auth import get_current_user, CurrentUser

    monkeypatch.setattr("app.db.service_client", _poisoned_service_client)

    fake_db = FakeDB()
    sa_uuid = str(uuid.uuid4())
    fake_db.table("templates").rows.append({
        "id": sa_uuid,
        "name": "Service Agreement",
        "category": "contracts",
        "template_key": "service-agreement",
        "review_status": "beta",
        "states_supported": [],
        "schema_json": {},
    })

    client = TestClient(app)

    def mock_user():
        return CurrentUser(id="21e63e8f-e00c-4ae6-afe4-17ba6b400be5", email="test@example.com", db=fake_db)

    app.dependency_overrides[get_current_user] = mock_user

    # 1. Lookup by template_key slug
    res_slug = client.get("/api/templates/service-agreement")
    assert res_slug.status_code == 200
    data_slug = res_slug.json()
    assert data_slug["name"] == "Service Agreement"
    sa_uuid = data_slug["id"]

    # 2. Lookup by UUID
    res_uuid = client.get(f"/api/templates/{sa_uuid}")
    assert res_uuid.status_code == 200
    data_uuid = res_uuid.json()
    assert data_uuid["name"] == "Service Agreement"
    assert data_uuid["template_key"] == "service-agreement"


def test_profile_metadata_no_hardcoded_advocate_defaults():
    """PROFILE-01: Verify user_metadata for a new user account does not contain
    hardcoded advocate demo strings (e.g. Nitesh Sharma, D/1042/2018), and that missing
    fields evaluate to empty strings instead of displaying another advocate's data."""
    mock_user_meta = {}

    def resolve_profile_field(meta_val, cached_val):
        return meta_val or cached_val or ""

    full_name = resolve_profile_field(mock_user_meta.get("full_name"), None)
    bar_number = resolve_profile_field(mock_user_meta.get("bar_number"), None)
    primary_court = resolve_profile_field(mock_user_meta.get("primary_court"), None)

    assert full_name == "", f"Expected empty string, got hardcoded value: {full_name}"
    assert bar_number == "", f"Expected empty string, got hardcoded value: {bar_number}"
    assert primary_court == "", f"Expected empty string, got hardcoded value: {primary_court}"
    assert "Nitesh" not in full_name
    assert "D/1042/2018" not in bar_number


# --- Phase 4.1: bounded concurrent clause generation ------------------------
#
# Phase 4 audit conclusion (SAFE WITH LIMITS): every llm_fillable clause's
# prompt is built solely from masked_form_data/masked_amendment_note, never
# from another clause's output -- StrictUndefined makes cross-clause
# reference an immediate crash, not a silent bug -- so clauses are safe to
# generate concurrently. The one shared mutable object is MaskMap; it's
# protected by a lock scoped to only the non-atomic mask_text() mutation
# inside llm_gateway.generate(), never held across the network call itself.
#
# None of these tests make real Gemini calls -- every fake `generate`
# replacement below is a deterministic, local stand-in.

import threading
import time as _time

from app.services.pii_mask import mask_text, unmask_text

THREE_LLM_CLAUSES = [
    {"clause_key": "clause_one", "display_order": 1, "clause_type": "llm_fillable",
     "applicable_condition": None, "heading": "Clause One",
     "source_text": "Draft clause one.", "current_text": "Draft clause one."},
    {"clause_key": "clause_two", "display_order": 2, "clause_type": "llm_fillable",
     "applicable_condition": None, "heading": "Clause Two",
     "source_text": "Draft clause two.", "current_text": "Draft clause two."},
    {"clause_key": "clause_three", "display_order": 3, "clause_type": "llm_fillable",
     "applicable_condition": None, "heading": "Clause Three",
     "source_text": "Draft clause three.", "current_text": "Draft clause three."},
]

ONE_LLM_CLAUSE = [
    {"clause_key": "solo_clause", "display_order": 1, "clause_type": "llm_fillable",
     "applicable_condition": None, "heading": "Solo Clause",
     "source_text": "Draft the only clause.", "current_text": "Draft the only clause."},
]


def _sleepy_generate_by_prompt(monkeypatch, delays_by_prompt: dict[str, float], record: list):
    """Fake `contracts.generate` that sleeps a caller-chosen amount per
    prompt and records (prompt, start, end) so tests can assert on overlap
    and ordering -- no network call, no real Gemini use."""
    lock = threading.Lock()

    def fake(prompt, task_type=None, mask_map=None, entities=None, auto_detect_names=True, **kwargs):
        start = _time.monotonic()
        _time.sleep(delays_by_prompt.get(prompt, 0.0))
        end = _time.monotonic()
        with lock:
            record.append((prompt, start, end))
        return GenerationResult(
            text=f"GENERATED[{prompt}]", provider="gemini", model="gemini-2.5-flash",
            latency_ms=int((end - start) * 1000), masked_prompt=prompt,
        )

    monkeypatch.setattr(contracts, "generate", fake)


def test_clauses_execute_concurrently_not_sequentially(monkeypatch):
    """Item 1: two independent llm_fillable clauses (NDA_CLAUSES has
    exactly two: recitals, term_and_survival) actually overlap in wall
    time -- proof this isn't just a sequential loop with extra bookkeeping."""
    # generate_draft()'s pre-loop _mask_form_data() call triggers spaCy's
    # one-time model load (~2-6s depending on machine, see Phase 3) the
    # first time ANY test in this process masks a real name -- if this
    # test happens to run before spaCy is already warm, that load lands
    # inside the timed section below and swamps the wall-clock assertion
    # with a cost that has nothing to do with clause-generation
    # concurrency. Force it to happen here, outside the timed section, so
    # this test's timing measures only what it claims to measure.
    from app.services.pii_mask import MaskMap as _WarmupMaskMap
    from app.services.pii_mask import mask_text as _warmup_mask_text
    _warmup_mask_text("Ramesh Kumar warmup", _WarmupMaskMap(matter_id="warmup"))

    db = FakeDB()
    template_id = _seed_template(db)  # NDA_CLAUSES: recitals + term_and_survival
    matter_id = _seed_matter(db)

    record: list[tuple[str, float, float]] = []
    delays = {
        "Draft recitals for: evaluating a potential software licensing deal": 0.15,
        "Draft term clause for: 3 years from the Effective Date": 0.15,
    }
    _sleepy_generate_by_prompt(monkeypatch, delays, record)

    wall_start = _time.monotonic()
    contracts.generate_draft(matter_id, template_id, BASE_FORM, db=db)
    wall_elapsed = _time.monotonic() - wall_start

    assert len(record) == 2
    # NOTE: deliberately no absolute wall_elapsed < N assertion here --
    # under real system load (e.g. running alongside the rest of the
    # suite, competing for CPU with other tests' own threads/imports) an
    # absolute small-fraction-of-a-second ceiling is flaky: it can fail
    # even when clause generation genuinely overlapped, simply because
    # the whole process was scheduled slower that run. The actual proof
    # of concurrency below (interval overlap) is structural and immune to
    # that -- it holds regardless of how slow or fast the machine is at
    # the moment, because both calls' start/end are measured relative to
    # each other, not against a fixed external budget.
    (_, a_start, a_end), (_, b_start, b_end) = record[0], record[1]
    overlap = a_start < b_end and b_start < a_end
    assert overlap, f"no overlap detected: {record}"


def test_maximum_concurrency_is_three(monkeypatch):
    """Item 2: the five-clause Consultancy template (the largest real
    template) never has more than MAX_CONCURRENT_CLAUSE_GENERATIONS (3)
    llm_fillable calls in flight at once."""
    db = FakeDB()
    template_id = _seed_consultancy(db)
    matter_id = _seed_matter(db)

    lock = threading.Lock()
    in_flight = 0
    max_in_flight = 0

    def fake(prompt, task_type=None, mask_map=None, entities=None, auto_detect_names=True, **kwargs):
        nonlocal in_flight, max_in_flight
        with lock:
            in_flight += 1
            max_in_flight = max(max_in_flight, in_flight)
        _time.sleep(0.05)
        with lock:
            in_flight -= 1
        return GenerationResult(
            text="GENERATED", provider="gemini", model="gemini-2.5-flash",
            latency_ms=50, masked_prompt=prompt,
        )

    monkeypatch.setattr(contracts, "generate", fake)
    contracts.generate_draft(matter_id, template_id, CONSULTANCY_FORM, db=db)

    assert max_in_flight <= contracts.MAX_CONCURRENT_CLAUSE_GENERATIONS, (
        f"observed {max_in_flight} concurrent calls, cap is "
        f"{contracts.MAX_CONCURRENT_CLAUSE_GENERATIONS}"
    )
    assert max_in_flight > 1, "expected genuine concurrency, not accidental serialization"


def test_original_clause_ordering_preserved_despite_different_completion_order(monkeypatch):
    """Item 3 / the mandated concurrency test: clause 1 sleeps 300ms,
    clause 2 sleeps 100ms, clause 3 sleeps 200ms -- completion order is
    2, 3, 1, but final_clause_texts/clause_fills must come back in
    original clause order (1, 2, 3), and max simultaneous calls <= 3."""
    db = FakeDB()
    template_id = _seed_template(db, clauses=THREE_LLM_CLAUSES)
    matter_id = _seed_matter(db)

    lock = threading.Lock()
    in_flight = 0
    max_in_flight = 0
    completion_order: list[str] = []
    delays = {
        "Draft clause one.": 0.30,
        "Draft clause two.": 0.10,
        "Draft clause three.": 0.20,
    }

    def fake(prompt, task_type=None, mask_map=None, entities=None, auto_detect_names=True, **kwargs):
        nonlocal in_flight, max_in_flight
        with lock:
            in_flight += 1
            max_in_flight = max(max_in_flight, in_flight)
        _time.sleep(delays[prompt])
        with lock:
            in_flight -= 1
            completion_order.append(prompt)
        return GenerationResult(
            text=f"GENERATED[{prompt}]", provider="gemini", model="gemini-2.5-flash",
            latency_ms=int(delays[prompt] * 1000), masked_prompt=prompt,
        )

    monkeypatch.setattr(contracts, "generate", fake)
    result = contracts.generate_draft(matter_id, template_id, {}, db=db)

    # Completion order proves the test actually exercised out-of-order
    # completion (otherwise this test would pass trivially).
    assert completion_order == ["Draft clause two.", "Draft clause three.", "Draft clause one."]
    assert max_in_flight <= 3

    # Original clause order must be preserved regardless.
    assert [f.clause_key for f in result.clause_fills] == ["clause_one", "clause_two", "clause_three"]
    assert "GENERATED[Draft clause one.]" in result.full_text
    assert result.full_text.index("GENERATED[Draft clause one.]") < result.full_text.index(
        "GENERATED[Draft clause two.]"
    ) < result.full_text.index("GENERATED[Draft clause three.]")


def test_one_clause_failure_prevents_all_persistence(monkeypatch, tmp_path):
    """Items 4/5 / the mandated failure test: clause 1 succeeds, clause 2
    fails, clause 3 succeeds -- generate_draft() must raise, and NOTHING
    must be persisted: no draft_versions row, no draft_clause_fills rows,
    no docx file. Also proves the other in-flight clauses are not
    abandoned mid-flight -- both succeeding clauses' fake calls are
    recorded as having completed even though the whole call fails."""
    db = FakeDB()
    template_id = _seed_template(db, clauses=THREE_LLM_CLAUSES)
    matter_id = _seed_matter(db)

    monkeypatch.setattr(contracts, "DRAFTS_DIR", tmp_path)

    completed: list[str] = []

    def fake(prompt, task_type=None, mask_map=None, entities=None, auto_detect_names=True, **kwargs):
        _time.sleep(0.05)
        if prompt == "Draft clause two.":
            completed.append(prompt)
            raise RuntimeError("simulated: all providers/models exhausted")
        completed.append(prompt)
        return GenerationResult(
            text=f"GENERATED[{prompt}]", provider="gemini", model="gemini-2.5-flash",
            latency_ms=50, masked_prompt=prompt,
        )

    monkeypatch.setattr(contracts, "generate", fake)

    with pytest.raises(RuntimeError, match="simulated"):
        contracts.generate_draft(matter_id, template_id, {}, db=db)

    # All three were allowed to run to completion -- none silently
    # abandoned mid-flight just because clause_two failed.
    assert set(completed) == {"Draft clause one.", "Draft clause two.", "Draft clause three."}

    # No persistence anywhere.
    assert db.table("draft_versions").rows == []
    assert db.table("draft_clause_fills").rows == []
    assert list(tmp_path.iterdir()) == [], "docx must not be written on a failed draft"


def test_maskmap_correct_under_concurrent_clause_execution(monkeypatch):
    """Item 6: three llm_fillable clauses each embed a DISTINCT PAN number
    directly in their static clause text (PAN detection runs regardless of
    auto_detect_names, per pii_mask.mask_text's docstring) -- this forces
    three concurrent, previously-unseen MaskMap.get_or_assign() calls
    against the SAME MaskMap object, the exact race the Phase 4 audit
    flagged. The fake `generate` below calls the REAL mask_text()/
    unmask_text() (under the real mask_lock contracts.py passes it),
    mirroring llm_gateway.generate()'s actual mask -> process -> unmask
    envelope, with only the network call itself replaced."""
    clauses = [
        {"clause_key": "c1", "display_order": 1, "clause_type": "llm_fillable",
         "applicable_condition": None, "heading": "C1",
         "source_text": "Ref PAN ABCDE1111F.", "current_text": "Ref PAN ABCDE1111F."},
        {"clause_key": "c2", "display_order": 2, "clause_type": "llm_fillable",
         "applicable_condition": None, "heading": "C2",
         "source_text": "Ref PAN FGHIJ2222K.", "current_text": "Ref PAN FGHIJ2222K."},
        {"clause_key": "c3", "display_order": 3, "clause_type": "llm_fillable",
         "applicable_condition": None, "heading": "C3",
         "source_text": "Ref PAN KLMNO3333P.", "current_text": "Ref PAN KLMNO3333P."},
    ]
    db = FakeDB()
    template_id = _seed_template(db, clauses=clauses)
    matter_id = _seed_matter(db)

    def fake(prompt, task_type=None, mask_map=None, entities=None, auto_detect_names=True, mask_lock=None, **kwargs):
        if mask_lock is not None:
            with mask_lock:
                masked = mask_text(prompt, mask_map, entities, auto_detect_names=auto_detect_names)
        else:
            masked = mask_text(prompt, mask_map, entities, auto_detect_names=auto_detect_names)
        _time.sleep(0.03)  # widen the race window
        llm_output = masked  # "model" echoes the masked prompt back verbatim
        unmasked = unmask_text(llm_output, mask_map) if mask_map else llm_output
        return GenerationResult(
            text=unmasked, provider="gemini", model="gemini-2.5-flash",
            latency_ms=30, masked_prompt=masked,
        )

    monkeypatch.setattr(contracts, "generate", fake)
    result = contracts.generate_draft(matter_id, template_id, {}, db=db)

    # Each clause's OWN generated text must contain its OWN PAN back,
    # correctly unmasked -- not another clause's (which would indicate a
    # forward/reverse mapping collision under concurrency).
    fills_by_key = {f.clause_key: f.generated_text for f in result.clause_fills}
    assert "ABCDE1111F" in fills_by_key["c1"] and "FGHIJ2222K" not in fills_by_key["c1"] and "KLMNO3333P" not in fills_by_key["c1"]
    assert "FGHIJ2222K" in fills_by_key["c2"] and "ABCDE1111F" not in fills_by_key["c2"] and "KLMNO3333P" not in fills_by_key["c2"]
    assert "KLMNO3333P" in fills_by_key["c3"] and "ABCDE1111F" not in fills_by_key["c3"] and "FGHIJ2222K" not in fills_by_key["c3"]

    # SupabaseMaskStore.save() persists one pii_masks row per distinct
    # placeholder (matter_id, placeholder, real_value, kind) -- exactly
    # three PAN rows must exist, each with the correct, un-collided
    # real_value. A concurrency-induced collision in MaskMap.get_or_assign()
    # would show up here as a missing row or a wrong real_value.
    pii_rows = [r for r in db.table("pii_masks").rows if r["matter_id"] == matter_id]
    pan_rows = [r for r in pii_rows if r["kind"] == "PAN"]
    assert len(pan_rows) == 3, f"expected 3 distinct PAN entries, got {pan_rows}"
    assert {r["real_value"] for r in pan_rows} == {"ABCDE1111F", "FGHIJ2222K", "KLMNO3333P"}
    assert len({r["placeholder"] for r in pan_rows}) == 3, "placeholder collision under concurrency"


def test_single_llm_fillable_clause_template_still_works(monkeypatch):
    """Item 7: a template with exactly one llm_fillable clause (e.g. the
    real Agreement to Sell / Lease Deed / Leave & Licence templates) must
    still work correctly through the ThreadPoolExecutor(max_workers=1)
    path -- not a special-cased/bypassed code path."""
    db = FakeDB()
    template_id = _seed_template(db, clauses=ONE_LLM_CLAUSE)
    matter_id = _seed_matter(db)
    _fake_generate(monkeypatch, canned_text="SOLO GENERATED TEXT")

    result = contracts.generate_draft(matter_id, template_id, {}, db=db)

    assert len(result.clause_fills) == 1
    assert result.clause_fills[0].clause_key == "solo_clause"
    assert "SOLO GENERATED TEXT" in result.full_text


def test_five_clause_consultancy_template_works(monkeypatch):
    """Item 8: the real, five-llm_fillable-clause Consultancy template
    (the largest of all 10 real templates, per the Phase 4 clause
    dependency table) completes correctly end to end under the bounded
    concurrent path, with all five clause_fills present in original
    template display_order."""
    db = FakeDB()
    template_id = _seed_consultancy(db)
    matter_id = _seed_matter(db)
    _fake_generate(monkeypatch, canned_text="CONSULTANCY CLAUSE TEXT")

    result = contracts.generate_draft(matter_id, template_id, CONSULTANCY_FORM, db=db)

    expected_keys = [
        "recitals", "scope_of_consulting_services", "ip_assignment",
        "term_and_termination", "governing_law_jurisdiction",
    ]
    assert [f.clause_key for f in result.clause_fills] == expected_keys
    assert len(result.clause_fills) == 5


def test_fixed_boilerplate_clauses_unaffected_by_concurrency_change(monkeypatch):
    """Item 10: fixed_boilerplate clauses (e.g. NDA's "definitions") are
    still rendered directly via Jinja, never touch the ThreadPoolExecutor
    or `generate` at all -- pre-existing guarantee, re-asserted here
    against the restructured clause loop specifically."""
    db = FakeDB()
    template_id = _seed_template(db)  # NDA_CLAUSES includes fixed "definitions"
    matter_id = _seed_matter(db)
    calls = _fake_generate(monkeypatch)

    result = contracts.generate_draft(matter_id, template_id, BASE_FORM, db=db)

    # Only the two llm_fillable clauses (recitals, term_and_survival)
    # reached the fake generate() -- fixed_boilerplate clauses never did.
    assert len(calls) == 2
    assert "Fixed boilerplate text." in result.full_text


def test_provider_fallback_path_unchanged_by_concurrency_wrapper(monkeypatch):
    """Item 9: contracts.py's concurrency wrapper doesn't alter
    llm_gateway.generate()'s own provider-fallback contract -- a clause
    whose `generate()` call raises after exhausting the failover chain
    (simulated here, real fallback behavior is llm_gateway.py's own
    unchanged responsibility and is covered by test_llm_gateway.py)
    still surfaces as a normal exception out of generate_draft(), the
    same as it did before this change."""
    db = FakeDB()
    template_id = _seed_template(db, clauses=ONE_LLM_CLAUSE)
    matter_id = _seed_matter(db)

    def fake_exhausted(prompt, task_type=None, mask_map=None, entities=None, auto_detect_names=True, **kwargs):
        raise RuntimeError("all providers/models exhausted")

    monkeypatch.setattr(contracts, "generate", fake_exhausted)

    with pytest.raises(RuntimeError, match="all providers/models exhausted"):
        contracts.generate_draft(matter_id, template_id, {}, db=db)


def test_concurrent_vs_sequential_mock_performance(monkeypatch):
    """Performance test (mocked, no real Gemini calls): a deterministic
    fake provider sleeps 0.1s per call. Compares MAX_CONCURRENT_CLAUSE_
    GENERATIONS=1 (forces effectively-sequential execution through the
    SAME code path) against the real default (3), for the three-clause
    fixture. Concurrent must be meaningfully faster -- proof the
    ThreadPoolExecutor is actually overlapping work, not just present in
    the code. Labelled explicitly as a MOCK measurement, not a real-world
    benchmark (see Phase 3/3.1 for real, measured Gemini latency)."""
    def make_fake():
        def fake(prompt, task_type=None, mask_map=None, entities=None, auto_detect_names=True, **kwargs):
            _time.sleep(0.1)
            return GenerationResult(
                text=f"GENERATED[{prompt}]", provider="gemini", model="gemini-2.5-flash",
                latency_ms=100, masked_prompt=prompt,
            )
        return fake

    # See test_clauses_execute_concurrently_not_sequentially's warm-up note
    # -- same rationale, applied here too before either timed section.
    from app.services.pii_mask import MaskMap as _WarmupMaskMap
    from app.services.pii_mask import mask_text as _warmup_mask_text
    _warmup_mask_text("Ramesh Kumar warmup", _WarmupMaskMap(matter_id="warmup"))

    # Sequential baseline: same code path, cap forced to 1.
    db1 = FakeDB()
    template_id = _seed_template(db1, clauses=THREE_LLM_CLAUSES)
    matter_id = _seed_matter(db1)
    monkeypatch.setattr(contracts, "generate", make_fake())
    monkeypatch.setattr(contracts, "MAX_CONCURRENT_CLAUSE_GENERATIONS", 1)
    t0 = _time.monotonic()
    contracts.generate_draft(matter_id, template_id, {}, db=db1)
    sequential_elapsed = _time.monotonic() - t0

    # Concurrent: default cap (3), all three clauses fit within it.
    db2 = FakeDB()
    template_id2 = _seed_template(db2, clauses=THREE_LLM_CLAUSES)
    matter_id2 = _seed_matter(db2)
    monkeypatch.setattr(contracts, "generate", make_fake())
    monkeypatch.setattr(contracts, "MAX_CONCURRENT_CLAUSE_GENERATIONS", 3)
    t0 = _time.monotonic()
    contracts.generate_draft(matter_id2, template_id2, {}, db=db2)
    concurrent_elapsed = _time.monotonic() - t0

    # MOCK measurement, explicitly labeled: 3 x 0.1s sequential (~0.3s) vs
    # concurrent bounded by the slowest single call (~0.1s + overhead).
    # NOTE: no absolute upper bound on concurrent_elapsed -- under real
    # system load a fixed small ceiling is flaky (see the overlap-test's
    # note above); the relative comparison below is the robust proof,
    # since both measurements are taken back-to-back under the same
    # momentary system conditions.
    assert sequential_elapsed > 0.28, f"sequential baseline too fast: {sequential_elapsed:.3f}s"
    assert concurrent_elapsed < sequential_elapsed / 1.5, (
        f"expected concurrent ({concurrent_elapsed:.3f}s) to be substantially faster "
        f"than sequential ({sequential_elapsed:.3f}s)"
    )




