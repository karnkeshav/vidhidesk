"""Document Composer (Sprint 3.6 Phase 2).

Assembles the latest APPROVED version of each of the 14 clause types
(clause_generator.py::CLAUSE_TYPES) into one ordered pleading. This module
contains NO legal reasoning and makes NO LLM calls — per the sprint's
explicit brief ("The composer itself must contain no legal reasoning. It
only assembles approved clauses.") it only:

  - preserves clause order (clause_generator.CLAUSE_TYPES, fixed)
  - preserves headings (clause_generator.CLAUSE_HEADINGS)
  - preserves numbering (sequential paragraph numbers assigned here, purely
    positional — never re-derived from clause content)
  - preserves citations (statute_refs / case_law_refs carried through
    verbatim from the clause row, never re-verified or re-graded here)
  - preserves clause order

A clause type with no approved version is listed in `missing_clauses`
rather than silently skipped or silently filled with placeholder prose —
an advocate composing a pleading with unreviewed clauses must see exactly
which sections are missing, not a document that reads as complete.
"""

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.services.clause_generator import CLAUSE_HEADINGS, CLAUSE_TYPES

# Same physical directory contracts.py's DRAFTS_DIR resolves to (computed
# independently rather than imported, so this module has no dependency on
# contracts.py — a composed pleading's export has nothing to do with the
# Contracts template/draft_versions pipeline, only the output location on
# disk happens to be shared, per this sprint's explicit "do not make
# Litigation use the Contracts draft export pipeline" instruction).
GENERATED_DRAFTS_DIR = Path(__file__).resolve().parent.parent.parent / "generated_drafts"

PLEADING_DISCLAIMER = "AI-generated draft for advocate review. Not legal advice."


class DocumentComposerError(ValueError):
    """Mirrors clause_generator.ClauseGeneratorError's plain
    ValueError-on-bad-input convention."""


def _next_version_no(matter_id: str, pleading_outline_id: str, db) -> int:
    rows = (
        db.table("litigation_pleading_drafts")
        .select("version_no")
        .eq("matter_id", matter_id)
        .eq("pleading_outline_id", pleading_outline_id)
        .order("version_no", desc=True)
        .limit(1)
        .execute()
        .data
    )
    return (rows[0]["version_no"] + 1) if rows else 1


def _latest_approved_by_type(matter_id: str, pleading_outline_id: str, db) -> dict[str, dict[str, Any]]:
    """Unlike clause_generator.latest_clauses_by_type (latest version
    regardless of review status), the composer must only ever pick a
    version an advocate has actually approved — the human-review gate this
    sprint's pipeline requires between clause generation and composition."""
    rows = (
        db.table("litigation_pleading_clauses")
        .select("*")
        .eq("matter_id", matter_id)
        .eq("pleading_outline_id", pleading_outline_id)
        .eq("review_status", "approved")
        .order("version_no", desc=True)
        .execute()
        .data
        or []
    )
    latest: dict[str, dict[str, Any]] = {}
    for row in rows:
        ct = row["clause_type"]
        if ct not in latest or row["version_no"] > latest[ct]["version_no"]:
            latest[ct] = row
    return latest


def compose_pleading(matter_id: str, pleading_outline_id: str, db) -> dict[str, Any]:
    """Assemble one composed pleading draft from whichever clause versions
    are currently approved. Always creates a new litigation_pleading_drafts
    row — recomposing after a clause changes is a new version, never an
    overwrite, same convention as every other artifact in this pipeline."""
    outline_rows = (
        db.table("litigation_pleading_outlines")
        .select("*")
        .eq("id", pleading_outline_id)
        .eq("matter_id", matter_id)
        .execute()
        .data
    )
    if not outline_rows:
        raise DocumentComposerError(
            f"Pleading outline {pleading_outline_id} not found for matter {matter_id}"
        )

    approved_by_type = _latest_approved_by_type(matter_id, pleading_outline_id, db)

    composed_sections: list[dict[str, Any]] = []
    clause_versions: list[dict[str, Any]] = []
    missing_clauses: list[str] = []

    for position, clause_type in enumerate(CLAUSE_TYPES, start=1):
        clause = approved_by_type.get(clause_type)
        if clause is None:
            missing_clauses.append(clause_type)
            continue
        content = clause.get("content") or {}
        composed_sections.append({
            "paragraph_no": position,
            "clause_type": clause_type,
            "heading": CLAUSE_HEADINGS[clause_type],
            "text": content.get("text", ""),
            "bullet_items": content.get("bullet_items"),
            "statute_refs": clause.get("statute_refs") or [],
            "case_law_refs": clause.get("case_law_refs") or [],
            "confidence": clause.get("confidence"),
        })
        clause_versions.append({
            "clause_type": clause_type,
            "clause_id": clause["id"],
            "version_no": clause["version_no"],
            "model_used": clause.get("model_used"),
            "prompt_version": clause.get("prompt_version"),
        })

    row = {
        "matter_id": matter_id,
        "pleading_outline_id": pleading_outline_id,
        "version_no": _next_version_no(matter_id, pleading_outline_id, db),
        "clause_versions": clause_versions,
        "composed_sections": composed_sections,
        "missing_clauses": missing_clauses,
        "created_at": datetime.now(timezone.utc).isoformat(),
    }
    inserted = db.table("litigation_pleading_drafts").insert(row).execute()
    return inserted.data[0] if inserted.data else row


def list_drafts(matter_id: str, pleading_outline_id: str, db) -> list[dict[str, Any]]:
    res = (
        db.table("litigation_pleading_drafts")
        .select("*")
        .eq("matter_id", matter_id)
        .eq("pleading_outline_id", pleading_outline_id)
        .order("version_no", desc=True)
        .execute()
    )
    return res.data or []


def get_draft(matter_id: str, draft_id: str, db) -> dict[str, Any] | None:
    """Single-draft fetch for export — filters by BOTH id and matter_id, so
    a draft belonging to a *different* matter (even one the same caller
    owns) never matches here, not just a different user's draft. `db` is
    expected to be the caller's RLS-scoped client (same convention as
    every other _get_*_or_404 helper in this codebase): a draft belonging
    to another user is already invisible to that query regardless of the
    explicit matter_id filter, no separate check needed — same
    no-ownership-probing-oracle posture as the rest of this codebase."""
    rows = (
        db.table("litigation_pleading_drafts")
        .select("*")
        .eq("id", draft_id)
        .eq("matter_id", matter_id)
        .limit(1)
        .execute()
        .data
    )
    return rows[0] if rows else None


def render_pleading_docx(draft: dict[str, Any]) -> Path:
    """Render an already-composed pleading draft's `composed_sections`
    into a downloadable .docx file. Deliberately NOT the docxtpl-skeleton-
    plus-Jinja-fill pattern contracts.py uses for Contracts templates —
    there is no fixed skeleton to fill here (a pleading's own "header"
    content, court/parties/case number, is itself the cause_title
    section's own composed text, produced upstream by
    clause_generator.py's deterministic builder). This is the smallest
    correct renderer for already-fully-assembled structured content:
    plain python-docx, one paragraph per line, in `composed_sections`'
    existing order — no legal reasoning, no LLM call, matching this
    module's own "assembly only" charter exactly.

    Always regenerates from `draft` rather than caching/reusing a
    previous file on disk — litigation_pleading_drafts rows are
    immutable (no UPDATE RLS policy, see migration 0018), so
    `composed_sections` for a given draft id never changes; regenerating
    is cheap (local CPU + file I/O only) and guarantees the file on disk
    can never drift from the row it was rendered from, at the cost of a
    small amount of redundant work on repeat downloads — a deliberately
    simple tradeoff, not an oversight.
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    banner = doc.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    banner_run = banner.add_run(PLEADING_DISCLAIMER)
    banner_run.bold = True
    banner_run.italic = True
    doc.add_paragraph()

    missing = draft.get("missing_clauses") or []
    if missing:
        warning = doc.add_paragraph()
        warning_run = warning.add_run(
            "INCOMPLETE DRAFT — the following sections have no advocate-"
            "approved clause version yet and are NOT included below: "
            + ", ".join(CLAUSE_HEADINGS.get(ct, ct) for ct in missing)
        )
        warning_run.bold = True
        doc.add_paragraph()

    for section in draft.get("composed_sections") or []:
        heading_para = doc.add_paragraph()
        heading_run = heading_para.add_run(
            f"{section.get('paragraph_no', '')}. {section.get('heading', '')}".strip()
        )
        heading_run.bold = True
        heading_run.font.size = Pt(12)

        bullet_items = section.get("bullet_items")
        if bullet_items:
            for item in bullet_items:
                doc.add_paragraph(str(item), style="List Bullet")
        else:
            for line in (section.get("text") or "").split("\n"):
                if line.strip():
                    doc.add_paragraph(line)
        doc.add_paragraph()

    GENERATED_DRAFTS_DIR.mkdir(parents=True, exist_ok=True)
    docx_path = GENERATED_DRAFTS_DIR / f"pleading_{draft['id']}.docx"
    doc.save(str(docx_path))
    return docx_path
