"""One-off builder for templates/rera/mortgage-deed.docx — the docxtpl
skeleton for the Mortgage Deed template (RERA & Real Estate, Phase 2B-1
backend).

Not run automatically; re-run manually (`python scripts/build_mortgage_deed_skeleton.py`)
only if the skeleton's fixed structure needs to change. The clause *content*
lives in Postgres (template_clauses, seeded by scripts/seed_mortgage_deed_template.py),
exactly the same split as every other Contracts/RERA template (see
build_nda_skeleton.py's own docstring for the full rationale — not
repeated here).
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUTPUT_PATH = Path(__file__).resolve().parent.parent.parent / "templates" / "rera" / "mortgage-deed.docx"


def build() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    banner = doc.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    banner_run = banner.add_run("{{ disclaimer_banner }}")
    banner_run.bold = True
    banner_run.italic = True

    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("MORTGAGE DEED")
    title_run.bold = True
    title_run.font.size = Pt(14)

    doc.add_paragraph()

    preamble = doc.add_paragraph()
    preamble.add_run(
        "THIS DEED OF MORTGAGE is made and executed at {{ property_state }} on this "
        "{{ execution_date }} by and between:"
    )

    doc.add_paragraph()

    mortgagor = doc.add_paragraph()
    mortgagor.add_run(
        "{{ mortgagor_name }}, residing/situated at {{ mortgagor_address }} "
        "(hereinafter referred to as the “MORTGAGOR”, which expression "
        "shall, unless repugnant to the context, include their heirs, "
        "legal representatives, successors, and assigns), of the FIRST "
        "PART;"
    )

    doc.add_paragraph().add_run("AND")

    mortgagee = doc.add_paragraph()
    mortgagee.add_run(
        "{{ mortgagee_name }}, residing/situated at {{ mortgagee_address }} "
        "(hereinafter referred to as the “MORTGAGEE”, which expression "
        "shall, unless repugnant to the context, include their heirs, "
        "legal representatives, successors, and assigns), of the SECOND "
        "PART."
    )

    doc.add_paragraph()

    # Single placeholder for the merged, ordered clause content — built at
    # render time as a docxtpl Subdocument by app/services/contracts.py,
    # exactly like every Contracts/RERA template. Must use the
    # paragraph-level {{p name}} tag (see build_nda_skeleton.py's comment
    # for why).
    doc.add_paragraph().add_run("{{p clauses_subdoc}}")

    doc.add_paragraph()
    doc.add_paragraph().add_run(
        "IN WITNESS WHEREOF, the MORTGAGOR and the MORTGAGEE have set "
        "their hands to this Deed of Mortgage on the day, month, and "
        "year first above written, in the presence of the following "
        "witnesses."
    )
    doc.add_paragraph()

    sig_mortgagor = doc.add_paragraph()
    sig_mortgagor.add_run("MORTGAGOR:").bold = True
    doc.add_paragraph("Signature: _______________________")
    doc.add_paragraph("Name: {{ mortgagor_name }}")
    doc.add_paragraph()

    sig_mortgagee = doc.add_paragraph()
    sig_mortgagee.add_run("MORTGAGEE:").bold = True
    doc.add_paragraph("Signature: _______________________")
    doc.add_paragraph("Name: {{ mortgagee_name }}")
    doc.add_paragraph()

    witnesses = doc.add_paragraph()
    witnesses.add_run("WITNESSES:").bold = True
    doc.add_paragraph("1. Signature: _______________________  Name: _______________  Address: _______________")
    doc.add_paragraph("2. Signature: _______________________  Name: _______________  Address: _______________")

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
