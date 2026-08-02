"""One-off builder for templates/contracts/service-agreement.docx — the
docxtpl skeleton for the Service Agreement template (Sprint 2 Deliverable
2, Batch 1).

Not run automatically; re-run manually
(`python scripts/build_service_agreement_skeleton.py`) only if the
skeleton's fixed structure needs to change. Clause content lives in
Postgres (template_clauses, seeded by scripts/seed_service_agreement_template.py).

Unlike NDA, this template has no variant concept — always asymmetric
Service Provider / Client, so party role labels are hardcoded literal
text here, not the `{{ party_a_role_label }}` mechanism NDA uses (that's
only populated when the template declares variant_field == "nda_variant";
see app/services/contracts.py::generate_draft).
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUTPUT_PATH = (
    Path(__file__).resolve().parent.parent.parent / "templates" / "contracts" / "service-agreement.docx"
)


def build() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    banner = doc.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    banner_run = banner.add_run("{{ disclaimer_banner }}")
    banner_run.bold = True
    banner_run.italic = True

    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("SERVICE AGREEMENT")
    title_run.bold = True
    title_run.font.size = Pt(14)

    doc.add_paragraph()

    preamble = doc.add_paragraph()
    preamble.add_run(
        "This Service Agreement (“Agreement”) is made and entered into as of "
        "{{ effective_date }} (“Effective Date”) by and between:"
    )

    doc.add_paragraph()

    party_a = doc.add_paragraph()
    party_a.add_run(
        "{{ party_a_name }}, {{ party_a_entity_type | an_or_a }} {{ party_a_entity_type }} having its "
        "registered / residential address at {{ party_a_address }} (hereinafter "
        "referred to as the “Service Provider”, which expression shall, unless "
        "repugnant to the context, include its successors and permitted assigns);"
    )

    doc.add_paragraph().add_run("AND")

    party_b = doc.add_paragraph()
    party_b.add_run(
        "{{ party_b_name }}, {{ party_b_entity_type | an_or_a }} {{ party_b_entity_type }} having its "
        "registered / residential address at {{ party_b_address }} (hereinafter "
        "referred to as the “Client”, which expression shall, unless repugnant "
        "to the context, include its successors and permitted assigns)."
    )

    doc.add_paragraph().add_run("(each a “Party” and collectively the “Parties”).")

    doc.add_paragraph()

    # Same {{p clauses_subdoc}} mechanism as NDA — see
    # docs/lessons_learned.md for why the plain {{ }} tag silently drops
    # the merged clause content instead of raising.
    doc.add_paragraph().add_run("{{p clauses_subdoc}}")

    doc.add_paragraph()
    doc.add_paragraph().add_run(
        "IN WITNESS WHEREOF, the Parties have executed this Agreement as of the "
        "Effective Date."
    )
    doc.add_paragraph()

    sig_a = doc.add_paragraph()
    sig_a.add_run("For and on behalf of {{ party_a_name }} (Service Provider):").bold = True
    doc.add_paragraph("Signature: _______________________")
    doc.add_paragraph("Name:")
    doc.add_paragraph("Designation:")
    doc.add_paragraph()

    sig_b = doc.add_paragraph()
    sig_b.add_run("For and on behalf of {{ party_b_name }} (Client):").bold = True
    doc.add_paragraph("Signature: _______________________")
    doc.add_paragraph("Name:")
    doc.add_paragraph("Designation:")

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
