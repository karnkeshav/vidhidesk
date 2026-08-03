"""Builder for templates/contracts/lease-deed.docx skeleton."""

from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUTPUT_PATH = (
    Path(__file__).resolve().parent.parent.parent / "templates" / "contracts" / "lease-deed.docx"
)


def build() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    banner = doc.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    banner_run = banner.add_run("{{ disclaimer_banner }}")
    banner_run.bold = True
    banner_run.italic = True

    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("LEASE DEED")
    title_run.bold = True
    title_run.font.size = Pt(14)

    doc.add_paragraph()

    preamble = doc.add_paragraph()
    preamble.add_run(
        "This Lease Deed (“Deed”) is made and executed at {{ state }} on this "
        "{{ effective_date }} by and between:"
    )

    doc.add_paragraph()

    party_a = doc.add_paragraph()
    party_a.add_run(
        "{{ party_a_name }}, {{ party_a_entity_type | an_or_a }} {{ party_a_entity_type }} having registered office / address at "
        "{{ party_a_address }} (hereinafter referred to as the “Lessor”, which expression shall include its legal heirs, "
        "successors and permitted assigns);"
    )

    doc.add_paragraph().add_run("AND")

    party_b = doc.add_paragraph()
    party_b.add_run(
        "{{ party_b_name }}, {{ party_b_entity_type | an_or_a }} {{ party_b_entity_type }} having registered office / address at "
        "{{ party_b_address }} (hereinafter referred to as the “Lessee”, which expression shall include its successors "
        "and permitted assigns)."
    )

    doc.add_paragraph().add_run("(The Lessor and Lessee are hereinafter individually referred to as “Party” and collectively as “Parties”).")

    doc.add_paragraph()

    # Subdocument insertion tag {{p clauses_subdoc}}
    doc.add_paragraph().add_run("{{p clauses_subdoc}}")

    doc.add_paragraph()
    doc.add_paragraph().add_run(
        "IN WITNESS WHEREOF, the Parties hereto have executed this Lease Deed on the day, month and year first written above."
    )
    doc.add_paragraph()

    sig_a = doc.add_paragraph()
    sig_a.add_run("LESSOR: {{ party_a_name }}").bold = True
    doc.add_paragraph("Signature: _______________________")
    doc.add_paragraph("Witness 1: ______________________")
    doc.add_paragraph()

    sig_b = doc.add_paragraph()
    sig_b.add_run("LESSEE: {{ party_b_name }}").bold = True
    doc.add_paragraph("Signature: _______________________")
    doc.add_paragraph("Witness 2: ______________________")

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
