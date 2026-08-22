"""One-off builder for templates/rera/relinquishment-deed.docx — the
docxtpl skeleton for the Relinquishment Deed template (RERA & Real
Estate, Phase 2B-1 backend).

Not run automatically; re-run manually (`python scripts/build_relinquishment_deed_skeleton.py`)
only if the skeleton's fixed structure needs to change. The clause *content*
lives in Postgres (template_clauses, seeded by scripts/seed_relinquishment_deed_template.py),
exactly the same split as every other Contracts/RERA template (see
build_nda_skeleton.py's own docstring for the full rationale — not
repeated here).
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUTPUT_PATH = Path(__file__).resolve().parent.parent.parent / "templates" / "rera" / "relinquishment-deed.docx"


def build() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    banner = doc.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    banner_run = banner.add_run("{{ disclaimer_banner }}")
    banner_run.bold = True
    banner_run.italic = True

    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("RELINQUISHMENT DEED")
    title_run.bold = True
    title_run.font.size = Pt(14)

    doc.add_paragraph()

    preamble = doc.add_paragraph()
    preamble.add_run(
        "THIS DEED OF RELINQUISHMENT is made and executed at "
        "{{ property_state }} on this {{ execution_date }} by and "
        "between:"
    )

    doc.add_paragraph()

    releasor = doc.add_paragraph()
    releasor.add_run(
        "{{ releasor_name }}, residing/situated at {{ releasor_address }} "
        "(hereinafter referred to as the “RELEASOR”, which expression "
        "shall, unless repugnant to the context, include their heirs, "
        "legal representatives, successors, and assigns), of the FIRST "
        "PART;"
    )

    doc.add_paragraph().add_run("AND")

    releasee = doc.add_paragraph()
    releasee.add_run(
        "{{ releasee_name }}, residing/situated at {{ releasee_address }} "
        "(hereinafter referred to as the “RELEASEE”, which expression "
        "shall, unless repugnant to the context, include their heirs, "
        "legal representatives, successors, and assigns), of the SECOND "
        "PART."
    )

    doc.add_paragraph()

    # Single placeholder for the merged, ordered clause content — built at
    # render time as a docxtpl Subdocument by app/services/contracts.py,
    # exactly like every Contracts/RERA template. Must use the
    # paragraph-level {{p name}} tag (see build_nda_skeleton.py's comment
    # for why).
    doc.add_paragraph().add_run("{{p clauses_subdoc}}")

    doc.add_paragraph()
    doc.add_paragraph().add_run(
        "IN WITNESS WHEREOF, the RELEASOR and the RELEASEE have set "
        "their hands to this Deed of Relinquishment on the day, month, "
        "and year first above written, in the presence of the following "
        "witnesses."
    )
    doc.add_paragraph()

    sig_releasor = doc.add_paragraph()
    sig_releasor.add_run("RELEASOR:").bold = True
    doc.add_paragraph("Signature: _______________________")
    doc.add_paragraph("Name: {{ releasor_name }}")
    doc.add_paragraph()

    sig_releasee = doc.add_paragraph()
    sig_releasee.add_run("RELEASEE:").bold = True
    doc.add_paragraph("Signature: _______________________")
    doc.add_paragraph("Name: {{ releasee_name }}")
    doc.add_paragraph()

    witnesses = doc.add_paragraph()
    witnesses.add_run("WITNESSES:").bold = True
    doc.add_paragraph("1. Signature: _______________________  Name: _______________  Address: _______________")
    doc.add_paragraph("2. Signature: _______________________  Name: _______________  Address: _______________")

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
