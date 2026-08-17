"""One-off builder for templates/rera/rera-complaint.docx — the docxtpl
skeleton for the RERA Complaint (Delay in Possession) template.

Not run automatically; re-run manually
(`python scripts/build_rera_complaint_skeleton.py`) only if the skeleton's
fixed structure needs to change. See build_nda_skeleton.py's docstring for
the full clause-vs-skeleton split rationale (identical here).
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUTPUT_PATH = Path(__file__).resolve().parent.parent.parent / "templates" / "rera" / "rera-complaint.docx"


def build() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    banner = doc.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    banner_run = banner.add_run("{{ disclaimer_banner }}")
    banner_run.bold = True
    banner_run.italic = True

    doc.add_paragraph()

    court = doc.add_paragraph()
    court.alignment = WD_ALIGN_PARAGRAPH.CENTER
    court.add_run(
        "BEFORE THE REAL ESTATE REGULATORY AUTHORITY, {{ jurisdiction_state | upper }}"
    ).bold = True

    doc.add_paragraph()

    case_no = doc.add_paragraph()
    case_no.alignment = WD_ALIGN_PARAGRAPH.CENTER
    case_no.add_run("Complaint No. _______________ of {{ agreement_date }}")

    doc.add_paragraph()

    in_matter = doc.add_paragraph()
    in_matter.add_run("IN THE MATTER OF:")
    doc.add_paragraph()

    complainant = doc.add_paragraph()
    complainant.add_run(
        "{{ complainant_name }}, residing at {{ complainant_address }}"
    )
    doc.add_paragraph().add_run("... COMPLAINANT")

    doc.add_paragraph().add_run("VERSUS")

    respondent = doc.add_paragraph()
    respondent.add_run(
        "{{ respondent_name }}, having its registered office at "
        "{{ respondent_address }}"
    )
    doc.add_paragraph().add_run("... RESPONDENT")

    doc.add_paragraph()

    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.add_run(
        "COMPLAINT UNDER SECTION 31 OF THE REAL ESTATE (REGULATION AND "
        "DEVELOPMENT) ACT, 2016"
    ).bold = True

    doc.add_paragraph()

    respectfully = doc.add_paragraph()
    respectfully.add_run(
        "MOST RESPECTFULLY SHOWETH:"
    )

    doc.add_paragraph()

    # Single placeholder for the merged, ordered clause content — same
    # docxtpl Subdocument pattern as every other template.
    doc.add_paragraph().add_run("{{p clauses_subdoc}}")

    doc.add_paragraph()

    place = doc.add_paragraph()
    place.add_run("Place: {{ jurisdiction_state }}")
    date = doc.add_paragraph()
    date.add_run("Date: {{ agreement_date }}")

    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.add_run("COMPLAINANT").bold = True
    doc.add_paragraph("Signature: _______________________")
    doc.add_paragraph("Name: {{ complainant_name }}")
    doc.add_paragraph()
    doc.add_paragraph("Through Advocate: _______________________")

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
