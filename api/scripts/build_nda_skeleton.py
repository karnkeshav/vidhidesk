"""One-off builder for templates/contracts/nda.docx — the docxtpl skeleton
for the NDA template (Sprint 2, Deliverable 1).

Not run automatically; re-run manually (`python scripts/build_nda_skeleton.py`)
only if the skeleton's fixed structure (header/party-block/signature) needs
to change. The clause *content* itself lives in Postgres (template_clauses,
seeded by scripts/seed_nda_template.py) and is reviewed there — this script
only builds the docx shell those clauses get rendered into.

Structure, per CLAUDE.md Hard Rule 2 (no hallucinated structure): this
skeleton carries the party/date/signature block via plain Jinja field
substitution. The operative clause text itself never lives in this file —
it's injected as a single merged Subdocument (`clauses_subdoc`) built at
render time by app/services/contracts.py from the ordered, variant-filtered
template_clauses rows. Adding/removing/reordering a clause is a Postgres
change, not a docx change.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUTPUT_PATH = Path(__file__).resolve().parent.parent.parent / "templates" / "contracts" / "nda.docx"


def build() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    banner = doc.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    banner_run = banner.add_run("{{ disclaimer_banner }}")
    banner_run.bold = True
    banner_run.italic = True

    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("NON-DISCLOSURE AGREEMENT")
    title_run.bold = True
    title_run.font.size = Pt(14)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("({{ nda_variant_label }})")

    doc.add_paragraph()

    preamble = doc.add_paragraph()
    preamble.add_run(
        "This Non-Disclosure Agreement (“Agreement”) is made and entered into "
        "as of {{ effective_date }} (“Effective Date”) by and between:"
    )

    doc.add_paragraph()

    party_a = doc.add_paragraph()
    party_a.add_run(
        "{{ party_a_name }}, {{ party_a_entity_type | an_or_a }} {{ party_a_entity_type }} having its registered / "
        "residential address at {{ party_a_address }} (hereinafter referred to as "
        "the “{{ party_a_role_label }}”, which expression shall, unless "
        "repugnant to the context, include its successors and permitted assigns);"
    )

    doc.add_paragraph().add_run("AND")

    party_b = doc.add_paragraph()
    party_b.add_run(
        "{{ party_b_name }}, {{ party_b_entity_type | an_or_a }} {{ party_b_entity_type }} having its registered / "
        "residential address at {{ party_b_address }} (hereinafter referred to as "
        "the “{{ party_b_role_label }}”, which expression shall, unless "
        "repugnant to the context, include its successors and permitted assigns)."
    )

    doc.add_paragraph().add_run(
        "(each a “Party” and collectively the “Parties”)."
    )

    doc.add_paragraph()

    # Single placeholder for the merged, ordered, variant-filtered clause
    # content — built at render time as a docxtpl Subdocument, not authored
    # here (see app/services/contracts.py). Must use docxtpl's paragraph-
    # level tag syntax `{{p name}}` (not `{{ name }}`) — a Subdocument's
    # value is raw multi-paragraph XML, which is only valid spliced in as
    # a sibling of this paragraph, not as text content inside it. A plain
    # `{{ }}` tag was tried first and silently produced an empty document:
    # the subdoc XML got dropped as invalid content inside a `<w:t>` run
    # rather than raising a visible error.
    doc.add_paragraph().add_run("{{p clauses_subdoc}}")

    doc.add_paragraph()
    doc.add_paragraph().add_run(
        "IN WITNESS WHEREOF, the Parties have executed this Agreement as of the "
        "Effective Date."
    )
    doc.add_paragraph()

    sig_a = doc.add_paragraph()
    sig_a.add_run("For and on behalf of {{ party_a_name }}:").bold = True
    doc.add_paragraph("Signature: _______________________")
    doc.add_paragraph("Name:")
    doc.add_paragraph("Designation:")
    doc.add_paragraph()

    sig_b = doc.add_paragraph()
    sig_b.add_run("For and on behalf of {{ party_b_name }}:").bold = True
    doc.add_paragraph("Signature: _______________________")
    doc.add_paragraph("Name:")
    doc.add_paragraph("Designation:")

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
