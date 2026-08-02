"""One-off builder for templates/contracts/employment.docx — the docxtpl
skeleton for the Employment Agreement template (Sprint 2 Deliverable 2,
Batch 4).

Not run automatically; re-run manually
(`python scripts/build_employment_skeleton.py`) only if the skeleton's
fixed structure needs to change. Clause content lives in Postgres
(template_clauses, seeded by scripts/seed_employment_template.py).

Unlike every prior template, Party B (the Employee) is always a natural
person — no party_b_entity_type field exists in the schema, so Party B's
block hardcodes "an Individual" rather than using the an_or_a filter
against a variable. Party A (the Employer) keeps the usual an_or_a
treatment since it can be any entity type.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUTPUT_PATH = (
    Path(__file__).resolve().parent.parent.parent / "templates" / "contracts" / "employment.docx"
)


def build() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    banner = doc.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    banner_run = banner.add_run("{{ disclaimer_banner }}")
    banner_run.bold = True
    banner_run.italic = True

    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("EMPLOYMENT AGREEMENT")
    title_run.bold = True
    title_run.font.size = Pt(14)

    doc.add_paragraph()

    preamble = doc.add_paragraph()
    preamble.add_run(
        "This Employment Agreement (“Agreement”) is made and entered into "
        "as of {{ effective_date }} (“Effective Date” / “Date of "
        "Joining”) by and between:"
    )

    doc.add_paragraph()

    party_a = doc.add_paragraph()
    party_a.add_run(
        "{{ party_a_name }}, {{ party_a_entity_type | an_or_a }} {{ party_a_entity_type }} having its "
        "registered / principal address at {{ party_a_address }} (hereinafter "
        "referred to as the “Employer”, which expression shall, unless "
        "repugnant to the context, include its successors and permitted assigns);"
    )

    doc.add_paragraph().add_run("AND")

    party_b = doc.add_paragraph()
    party_b.add_run(
        "{{ party_b_name }}, an Individual residing at {{ party_b_address }} "
        "(hereinafter referred to as the “Employee”)."
    )

    doc.add_paragraph().add_run("(each a “Party” and collectively the “Parties”).")

    doc.add_paragraph()

    # {{p clauses_subdoc}} — not {{ clauses_subdoc }} — see
    # docs/lessons_learned.md for why the plain tag silently drops the
    # merged clause content instead of raising.
    doc.add_paragraph().add_run("{{p clauses_subdoc}}")

    doc.add_paragraph()
    doc.add_paragraph().add_run(
        "IN WITNESS WHEREOF, the Parties have executed this Agreement as of "
        "the Effective Date."
    )
    doc.add_paragraph()

    sig_a = doc.add_paragraph()
    sig_a.add_run("For and on behalf of {{ party_a_name }} (Employer):").bold = True
    doc.add_paragraph("Signature: _______________________")
    doc.add_paragraph("Name:")
    doc.add_paragraph("Designation:")
    doc.add_paragraph()

    sig_b = doc.add_paragraph()
    sig_b.add_run("{{ party_b_name }} (Employee):").bold = True
    doc.add_paragraph("Signature: _______________________")

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
